import docx
from docx.shared import Pt, RGBColor
import re
import sys

md_path = "Presentacion_Auditores_IFRS16.md"
docx_path = "Presentacion_Auditores_IFRS16.docx"

try:
    with open(md_path, "r", encoding="utf-8") as f:
        content = f.read()
except Exception as e:
    sys.exit()

doc = docx.Document()

# Styles
for style in doc.styles:
    if hasattr(style, 'font'):
        style.font.name = 'Calibri'

title = doc.add_heading(level=0)
title_run = title.add_run("Validación Técnica del Motor Financiero IFRS 16")
title_run.font.color.rgb = RGBColor(0, 130, 155) # Mundo 16 Blue

doc.add_paragraph("Documento dirigido a Auditores Externos y Revisores de Cumplimiento Normativo", style='Subtitle')

lines = content.split('\n')
skip_next = False

for line in lines:
    line = line.strip()
    if not line or line.startswith('---') or line.startswith('# Validación') or line.startswith('**Documento'):
        continue
        
    if line.startswith('## '):
        h = doc.add_heading(level=1)
        # remove bold from heading string
        t = line[3:].replace('**', '')
        r = h.add_run(t)
        r.font.color.rgb = RGBColor(0,0,0)
    elif line.startswith('### '):
        doc.add_heading(line[4:].replace('**', ''), level=2)
    elif line.startswith('- '):
        p = doc.add_paragraph(style='List Bullet')
        # handle bold inside
        parts = re.split(r'(\*\*.*?\*\*)', line[2:])
        for p_t in parts:
            if p_t.startswith('**') and p_t.endswith('**'):
                p.add_run(p_t[2:-2]).bold = True
            else:
                p.add_run(p_t.replace('*', ''))
    elif line.startswith('>'):
        p = doc.add_paragraph()
        run = p.add_run(line[1:].strip().replace('**', ''))
        run.italic = True
    else:
        p = doc.add_paragraph()
        parts = re.split(r'(\*\*.*?\*\*)', line)
        for p_t in parts:
            if p_t.startswith('**') and p_t.endswith('**'):
                p.add_run(p_t[2:-2]).bold = True
            else:
                p.add_run(p_t.replace('*', ''))

doc.save(docx_path)
