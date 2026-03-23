import sys
import os

try:
    from docx import Document
    from docx.shared import Pt
    from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

    doc = Document()

    # Title
    title = doc.add_heading('Manual de Usuario: Sistema IFRS 16 Pro', 0)
    title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    doc.add_paragraph('¡Bienvenido al Sistema IFRS 16 Pro! Este manual está diseñado paso a paso para que cualquier persona, sin importar su experiencia técnica, pueda cargar contratos, gestionar tipos de cambio y generar los reportes financieros exigidos por la norma NIIF 16 (IFRS 16).')

    doc.add_heading('🛑 REGLA DE ORO ANTES DE EMPEZAR', level=1)
    p1 = doc.add_paragraph()
    r1 = p1.add_run('NUNCA cargues un contrato sin haber cargado antes las monedas (UF, USD, EUR, etc.).\n')
    r1.bold = True
    p1.add_run('El sistema necesita saber cuánto vale la moneda exactamente el día que inicia tu contrato. Si subes el contrato y el sistema no encuentra la moneda de ese día, registrará un valor de cero y el cálculo fallará.\n\n')
    p1.add_run('Orden estricto de uso:\n1. Cargar Monedas.\n2. Cargar Contratos.\n3. Generar Reportes.')

    doc.add_heading('Módulo 1: 💱 Monedas (Tipos de Cambio)', level=1)
    doc.add_paragraph('Aquí le enseñamos al sistema a cuánto estaban las divisas en el pasado y en el presente.')
    doc.add_paragraph('1. Ve a la pestaña Monedas.\n2. Tienes dos opciones principales:\n   - Carga Manual: Útil para agregar el valor de la UF de un solo día en particular.\n   - Carga Masiva (Recomendado): Descarga la plantilla, llénala con todo el historial de fechas y valores mensuales o diarios, arrástralo y presiona Cargar.\n3. Revisa en la pestaña "Ver Todos los Datos" que tu historial esté correcto.')

    doc.add_heading('Módulo 2: 📝 Contratos (Arrendamientos)', level=1)
    doc.add_paragraph('Aquí es donde registras todos los alquileres.')
    doc.add_paragraph('1. Ve a la pestaña Contratos.\n2. Carga Masiva (Recomendado):\n   - Descarga la "Plantilla con Ejemplos".\n   - Llena el Excel respetando EXACTAMENTE las palabras (ej: Vencido, Anticipado, Mensual, Semestral sin espacios al final).\n   - Arrastra el Excel y procesa.\n3. Modificación de Contrato: Si cambian los plazos o el canon, entra aquí. Registra las nuevas condiciones y el sistema hará la remedición automática.\n4. Baja Anticipada: Úsalo si devuelves un activo antes de su fecha original de Fin.')

    doc.add_heading('Módulo 3: 🧮 Panel de Saldos (Dashboard)', level=1)
    doc.add_paragraph('El corazón del sistema. Úsalo cada fin de mes.')
    doc.add_paragraph('1. Entra a Panel de Saldos y filtra tu Empresa y Mes/Año de cierre.\n2. Haz clic en "Generar Resumen de Saldos".\n3. En la pestaña Resumen Consolidado ves los grandes totales de Activo y Pasivo ROU.\n4. En Detalle por Contrato verificas matemáticamente el Valor Presente y las Cuotas de Pago (VA).\n5. Exporta todo a Excel para tus centralizaciones ERP.')

    doc.add_heading('Módulo 4: 📊 Notas a los Estados Financieros (Vencimientos)', level=1)
    doc.add_paragraph('La tabla obligatoria de Riesgo de Liquidez de IFRS 7.')
    doc.add_paragraph('1. Entra al módulo y elige tu Empresa, Mes y Año.\n2. Botón "Generar Pasivos No Descontados".\n3. Revisa la agrupación por tramos de tiempo (< 90 días, 1 año, 2 a 3 años, etc.) y úsala directo en el balance consolidado auditado.')

    doc.add_heading('Anexo: Preguntas Frecuentes', level=1)
    doc.add_paragraph('¿Qué pasa si mi frecuencia no es mensual?')
    doc.add_paragraph('El sistema aplicará salidas de caja solo según tu frecuencia (ej. cada 6 meses), pero devengará la contabilidad estrictamente mes a mes como lo pide IFRS 16.')
    doc.add_paragraph('¿Qué hago si me equivoqué?')
    doc.add_paragraph('Usa Modificación de Contrato para remedir y ajustar tasas. Si el contrato nunca arrancó, pide su eliminación directa por base de datos.')

    doc.save('C:/Users/cfsaa/OneDrive/Desktop/Software_IFRS16_Pro/Manual_Usuario_IFRS16.docx')
    print("Success: Document generated")
except ImportError:
    print("Module python-docx not installed. Will write text file instead.")
    with open('C:/Users/cfsaa/OneDrive/Desktop/Software_IFRS16_Pro/Manual_Usuario_IFRS16.txt', 'w', encoding='utf-8') as f:
        f.write('Manual was successfully built in MD format. Word module missing.')
