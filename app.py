import streamlit as st
import pandas as pd
from datetime import date
import datetime
from dateutil.relativedelta import relativedelta
import io
from db import *
from core import *
st.set_page_config(page_title="Mundo 16", layout="wide")

# -- INYECCIÓN AVANZADA DE ESTILOS CORPORATIVOS MUNDO --
st.markdown("""
<style>
    /* Importar fuente oficial corporativa de Google Fonts (Montserrat) */
    @import url('https://fonts.googleapis.com/css2?family=Montserrat:wght@400;600;800;900&display=swap');

    html, body, [class*="css"], .stMarkdown p, .stDataFrame {
        font-family: 'Montserrat', sans-serif !important;
    }
    
    /* Encabezados gruesos estilo "El Match Perfecto" */
    h1, h2, h3 {
        color: #00829B !important;
        font-weight: 800 !important;
        text-transform: uppercase !important;
        letter-spacing: -0.5px;
    }
    
    /* Escalando los tamaños para no opacar el contenido */
    h1 { font-size: 1.8rem !important; }
    h2 { font-size: 1.4rem !important; }
    h3 { font-size: 1.15rem !important; }
    
    /* Decoración de Títulos */
    h1, h2 {
        border-bottom: 5px solid #FFCE00 !important;
        padding-bottom: 5px !important;
        margin-bottom: 30px !important;
        display: inline-block;
    }

    /* Botones Redondeados (Píldoras) estilo "Fibra + TV" o "Sucursal Virtual" */
    .stButton > button {
        border-radius: 50px !important;
        font-weight: 700 !important;
        padding: 0.5rem 2rem !important;
        text-transform: uppercase;
        letter-spacing: 0.5px;
        transition: all 0.3s ease !important;
    }
    
    .stButton > button[kind="primary"] {
        background-color: #00829B !important;
        color: white !important;
        border: none !important;
    }
    
    .stButton > button[kind="primary"]:hover {
        background-color: #00607A !important;
        transform: translateY(-2px);
        box-shadow: 0 6px 15px rgba(0, 130, 155, 0.4) !important;
    }

    /* Estilización de las Pestañas (Tabs) como Toggles estilo "2 MUNDO / 1 MUNDO" */
    .stTabs [data-baseweb="tab-list"] {
        background-color: #FFFFFF;
        border-radius: 50px;
        padding: 5px;
        gap: 5px;
        display: inline-flex;
        border: 2px solid #EBEBEB;
        box-shadow: 0 4px 10px rgba(0,0,0,0.03);
    }

    .stTabs [data-baseweb="tab"] {
        border-radius: 50px !important;
        padding: 6px 16px !important;
        font-size: 0.85rem !important;
        font-weight: 800 !important;
        border: none !important;
        background-color: transparent !important;
        color: #B0B0B0 !important;
    }

    .stTabs [aria-selected="true"] {
        background-color: #00829B !important;
        color: #FFFFFF !important;
        box-shadow: 0 4px 12px rgba(0, 130, 155, 0.3) !important;
    }

    /* Ocultar barra inferior nativa de selección de Streamlit Tab */
    .stTabs [data-baseweb="tab-highlight"] {
        display: none !important;
    }

    /* Simulación de Top Bar (Banner Oscuro) */
    header[data-testid="stHeader"] {
        border-top: 15px solid #00607A !important;
    }
    
    header[data-testid="stHeader"]::after {
        content: "";
        display: block;
        height: 6px;
        width: 100%;
        background-color: #FFCE00;
    }
    
    /* Contenedores con efecto Tarjeta (Cards) */
    [data-testid="stVerticalBlock"] [data-testid="stVerticalBlock"] {
        background-color: #FFFFFF;
        border-radius: 15px;
        padding: 20px;
        box-shadow: 0 10px 30px rgba(0,0,0,0.04);
        border: 1px solid #F0F0F0;
    }
</style>
""", unsafe_allow_html=True)

if 'success_msg' in st.session_state:
    st.success(st.session_state.success_msg)
    del st.session_state.success_msg

# --- KILL SWITCH (Domingo 31 Mayo 2026) ---
FECHA_EXPIRACION = date(2026, 5, 31)
is_local = False
try:
    host = st.context.headers.get("Host", "")
    if "localhost" in host or "127.0.0.1" in host:
        is_local = True
except:
    pass

if not is_local and date.today() > FECHA_EXPIRACION:
    st.error("Acceso Denegado: La versión de prueba de Mundo 16 ha expirado. Contacte al administrador para soporte.")
    st.stop()


MESES_LISTA = ["Enero", "Febrero", "Marzo", "Abril", "Mayo", "Junio", "Julio", "Agosto", "Septiembre", "Octubre", "Noviembre", "Diciembre"]

if 'auth' not in st.session_state: st.session_state.auth = False
EMPRESAS_LISTA = obtener_parametros('EMPRESA')
CLASES_ACTIVO = obtener_parametros('CLASE_ACTIVO')

# --- FUNCIONES DE APOYO CONTABLE ---

def add_asiento(lista, emp, cod1, transaccion, n_cta, cuenta, debe, haber):
    # Forzar que solo exista Debe o Haber por línea. Si vienen ambos, se dividen.
    if abs(debe) > 0.01:
        lista.append({"Empresa": emp, "Cod1": cod1, "Transacción": transaccion, "N° Cuenta": str(n_cta), "Cuenta": str(cuenta), "Tipo": "Debe", "Debe": round(debe,0), "Haber": 0})
    if abs(haber) > 0.01:
        lista.append({"Empresa": emp, "Cod1": cod1, "Transacción": transaccion, "N° Cuenta": str(n_cta), "Cuenta": str(cuenta), "Tipo": "Haber", "Debe": 0, "Haber": round(haber,0)})

def obtener_motor_financiero(c, rems=None):
    if 'motor_cache' not in st.session_state:
        st.session_state.motor_cache = {}
    
    cid = c['Codigo_Interno']
    hash_c = f"{c['Estado']}_{c['Canon']}_{c['Tasa']}_{c['Plazo']}_{c['Inicio']}_{c['Fin']}_{c.get('Fecha_Baja', '')}_{len(rems) if rems else 0}"
    
    if cid in st.session_state.motor_cache:
        cached_hash, tab, vp, rou = st.session_state.motor_cache[cid]
        if cached_hash == hash_c:
            return tab, vp, rou
            
    tab, vp, rou = motor_financiero_v20(c, rems)
    st.session_state.motor_cache[cid] = (hash_c, tab, vp, rou)
    return tab, vp, rou


# --- MÓDULOS ---

def modulo_asientos():
    st.header("🧾 Registros Contables")
    c1, c2, c3 = st.columns(3)
    emp_sel = c1.selectbox("Empresa", ["Todas"] + EMPRESAS_LISTA, key="as_emp")
    m_nom = c2.selectbox("Mes", MESES_LISTA, key="as_m")
    a = c3.number_input("Año", value=date.today().year, key="as_a")
    
    if st.button("Generar Asientos", type="primary"):
        m_idx = MESES_LISTA.index(m_nom) + 1
        f_act = pd.to_datetime(date(a, m_idx, 1)) + relativedelta(day=31)
        f_ant = f_act - relativedelta(months=1, day=31)
        
        detalles = []
        lista_c = cargar_contratos()
        from db import cargar_remediciones_todas_agrupadas
        rems_grupos = cargar_remediciones_todas_agrupadas()
        
        if emp_sel != "Todas": lista_c = [c for c in lista_c if c['Empresa'] == emp_sel]
        
        cta_map = {
            'ROU': (obtener_parametros('CUENTA_ROU_NUM')[0] if obtener_parametros('CUENTA_ROU_NUM') else '1401', obtener_parametros('CUENTA_ROU_NOM')[0] if obtener_parametros('CUENTA_ROU_NOM') else 'Activo ROU'),
            'Pasivo': (obtener_parametros('CUENTA_PASIVO_NUM')[0] if obtener_parametros('CUENTA_PASIVO_NUM') else '2101', obtener_parametros('CUENTA_PASIVO_NOM')[0] if obtener_parametros('CUENTA_PASIVO_NOM') else 'Pasivo IFRS 16'),
            'Ajuste': (obtener_parametros('CUENTA_BCO_AJUSTE_NUM')[0] if obtener_parametros('CUENTA_BCO_AJUSTE_NUM') else '1101', obtener_parametros('CUENTA_BCO_AJUSTE_NOM')[0] if obtener_parametros('CUENTA_BCO_AJUSTE_NOM') else 'Banco Ajustes'),
            'Amort': (obtener_parametros('CUENTA_GASTO_AMORT_NUM')[0] if obtener_parametros('CUENTA_GASTO_AMORT_NUM') else '4101', obtener_parametros('CUENTA_GASTO_AMORT_NOM')[0] if obtener_parametros('CUENTA_GASTO_AMORT_NOM') else 'Gasto Amort'),
            'AmortAcum': (obtener_parametros('CUENTA_AMORT_ACUM_NUM')[0] if obtener_parametros('CUENTA_AMORT_ACUM_NUM') else '1402', obtener_parametros('CUENTA_AMORT_ACUM_NOM')[0] if obtener_parametros('CUENTA_AMORT_ACUM_NOM') else 'Amort Acumulada'),
            'Interes': (obtener_parametros('CUENTA_GASTO_INT_NUM')[0] if obtener_parametros('CUENTA_GASTO_INT_NUM') else '4201', obtener_parametros('CUENTA_GASTO_INT_NOM')[0] if obtener_parametros('CUENTA_GASTO_INT_NOM') else 'Gasto Interés'),
            'Banco': (obtener_parametros('CUENTA_BANCO_PAGO_NUM')[0] if obtener_parametros('CUENTA_BANCO_PAGO_NUM') else '1102', obtener_parametros('CUENTA_BANCO_PAGO_NOM')[0] if obtener_parametros('CUENTA_BANCO_PAGO_NOM') else 'Banco Efectivo'),
            'Perdida': (obtener_parametros('CUENTA_PERDIDA_TC_NUM')[0] if obtener_parametros('CUENTA_PERDIDA_TC_NUM') else '4301', obtener_parametros('CUENTA_PERDIDA_TC_NOM')[0] if obtener_parametros('CUENTA_PERDIDA_TC_NOM') else 'Pérdida TC'),
            'Ganancia': (obtener_parametros('CUENTA_GANANCIA_TC_NUM')[0] if obtener_parametros('CUENTA_GANANCIA_TC_NUM') else '4302', obtener_parametros('CUENTA_GANANCIA_TC_NOM')[0] if obtener_parametros('CUENTA_GANANCIA_TC_NOM') else 'Ganancia TC')
        }
        
        for c in lista_c:
            f_ini = pd.to_datetime(c['Inicio'])
            if f_act < f_ini.replace(day=1): continue
            
            tab, vp, rou = obtener_motor_financiero(c, rems=rems_grupos.get(c['Codigo_Interno'], []))
            if tab.empty or 'Fecha' not in tab.columns: continue
            
            # 1. Asiento de Reconocimiento Inicial
            if f_ini.month == m_idx and f_ini.year == a:
                tc_ini = float(c['Valor_Moneda_Inicio']) if c['Valor_Moneda_Inicio'] > 0 else 1.0
                t_rec = "1. Reconocimiento Inicial"
                add_asiento(detalles, c['Empresa'], c['Codigo_Interno'], t_rec, *cta_map['ROU'], rou * tc_ini, 0)
                add_asiento(detalles, c['Empresa'], c['Codigo_Interno'], t_rec, *cta_map['Pasivo'], 0, vp * tc_ini)
                diff = (rou - vp) * tc_ini
                if diff > 0:
                    add_asiento(detalles, c['Empresa'], c['Codigo_Interno'], t_rec, *cta_map['Ajuste'], 0, diff)
                elif diff < 0:
                    add_asiento(detalles, c['Empresa'], c['Codigo_Interno'], t_rec, *cta_map['Ajuste'], abs(diff), 0)
    
            # Asientos Mensuales (Amortización, Intereses, Pagos y Reajuste)
            fila = tab[(tab['Fecha'].dt.month == m_idx) & (tab['Fecha'].dt.year == a)]
            if not fila.empty:
                it = fila.iloc[0]
                tc_act = obtener_tc_cache(c['Moneda'], f_act)
                
                # Ajuste crucial: Si es la primera cuota (Mes 1), la UF anterior es la de inicio del contrato
                if it['Mes'] == 1:
                    tc_ant = float(c['Valor_Moneda_Inicio']) if c.get('Valor_Moneda_Inicio') and float(c['Valor_Moneda_Inicio']) > 0 else 1.0
                else:
                    tc_ant = obtener_tc_cache(c['Moneda'], f_ant)
                    
                tc_ini = float(c['Valor_Moneda_Inicio']) if c.get('Valor_Moneda_Inicio') and float(c['Valor_Moneda_Inicio']) > 0 else 1.0
                ratio_act = tc_act
                ratio_ant = tc_ant
                
                t_amo = "2. Amortización"
                # ROU es activo no monetario a costo histórico (NIC 21). Su amortización se ancla irrevocablemente al TC inicial.
                add_asiento(detalles, c['Empresa'], c['Codigo_Interno'], t_amo, *cta_map['Amort'], it['Dep_Orig'] * tc_ini, 0)
                add_asiento(detalles, c['Empresa'], c['Codigo_Interno'], t_amo, *cta_map['AmortAcum'], 0, it['Dep_Orig'] * tc_ini)
                
                t_pag = "3. Pago de Arriendo ROU"
                add_asiento(detalles, c['Empresa'], c['Codigo_Interno'], t_pag, *cta_map['Pasivo'], it['Pago_Orig'] * ratio_act, 0)
                add_asiento(detalles, c['Empresa'], c['Codigo_Interno'], t_pag, *cta_map['Banco'], 0, it['Pago_Orig'] * ratio_act)
                
                t_int = "4. Intereses"
                add_asiento(detalles, c['Empresa'], c['Codigo_Interno'], t_int, *cta_map['Interes'], it['Int_Orig'] * ratio_act, 0)
                add_asiento(detalles, c['Empresa'], c['Codigo_Interno'], t_int, *cta_map['Pasivo'], 0, it['Int_Orig'] * ratio_act)
                
                t_tc = "5. Diferencia de Cambio"
                reajuste = (it['S_Fin_Orig'] * ratio_act) - (it['S_Fin_Orig'] * ratio_ant)
                if abs(reajuste) > 0.1:
                    if reajuste > 0: 
                        add_asiento(detalles, c['Empresa'], c['Codigo_Interno'], t_tc, *cta_map['Perdida'], reajuste, 0)
                        add_asiento(detalles, c['Empresa'], c['Codigo_Interno'], t_tc, *cta_map['Pasivo'], 0, reajuste)
                    else:
                        add_asiento(detalles, c['Empresa'], c['Codigo_Interno'], t_tc, *cta_map['Pasivo'], abs(reajuste), 0)
                        add_asiento(detalles, c['Empresa'], c['Codigo_Interno'], t_tc, *cta_map['Ganancia'], 0, abs(reajuste))
    
            # Asiento de Baja Definitiva o Remedición Manual
            paso_baja_manual = False
            if c.get('Fecha_Baja') and c['Estado'] in ['Baja', 'Remedido']:
                f_baja = pd.to_datetime(c['Fecha_Baja'])
                if f_baja.month == m_idx and f_baja.year == a:
                    paso_baja_manual = True
                    # Calcular saldos al momento del cese/remedición
                    pasado = tab[tab['Fecha'] <= f_baja]
                    if not pasado.empty:
                        tc_baja = obtener_tc_cache(c['Moneda'], f_baja)
                        s_fin_pasivo = pasado.iloc[-1]['S_Fin_Orig'] * tc_baja
                        amort_acum = pasado['Dep_Orig'].sum() * tc_baja
                        s_fin_rou = (rou * tc_baja) - amort_acum
                        
                        if s_fin_pasivo > 0.01 or abs(rou * tc_baja) > 0.01:
                            if c['Estado'] == 'Remedido':
                                t_baja = "7. Cierre por Remedición"
                                # Reversar Pasivo
                                if s_fin_pasivo > 0:
                                    add_asiento(detalles, c['Empresa'], c['Codigo_Interno'], t_baja, *cta_map['Pasivo'], s_fin_pasivo, 0)
                                
                                # Reversar ROU y Amort
                                r_neto = (rou * tc_baja)
                                add_asiento(detalles, c['Empresa'], c['Codigo_Interno'], t_baja, *cta_map['ROU'], 0, r_neto)
                                add_asiento(detalles, c['Empresa'], c['Codigo_Interno'], t_baja, *cta_map['AmortAcum'], amort_acum, 0)
                                
                                # Diferencia a Ajuste (SIN P&L)
                                dif_baja = s_fin_pasivo - s_fin_rou
                                if dif_baja > 0:
                                    add_asiento(detalles, c['Empresa'], c['Codigo_Interno'], t_baja, *cta_map['Ajuste'], 0, dif_baja)
                                elif dif_baja < 0:
                                    add_asiento(detalles, c['Empresa'], c['Codigo_Interno'], t_baja, *cta_map['Ajuste'], abs(dif_baja), 0)
                            else:
                                t_baja = "6. Baja Definitiva de Contrato"
                                if s_fin_pasivo > 0:
                                    add_asiento(detalles, c['Empresa'], c['Codigo_Interno'], t_baja, *cta_map['Pasivo'], s_fin_pasivo, 0)
                                r_neto = (rou * tc_baja)
                                add_asiento(detalles, c['Empresa'], c['Codigo_Interno'], t_baja, *cta_map['ROU'], 0, r_neto)
                                add_asiento(detalles, c['Empresa'], c['Codigo_Interno'], t_baja, *cta_map['AmortAcum'], amort_acum, 0)
                                dif_baja = s_fin_pasivo - s_fin_rou
                                if dif_baja > 0:
                                    add_asiento(detalles, c['Empresa'], c['Codigo_Interno'], t_baja, *cta_map['Ganancia'], 0, dif_baja)
                                elif dif_baja < 0:
                                    add_asiento(detalles, c['Empresa'], c['Codigo_Interno'], t_baja, *cta_map['Perdida'], abs(dif_baja), 0)
    
            # Verificar si hubo una Terminación Parcial (Reducción) en el mismo mes que anule la Baja Natural
            paso_terminacion_parcial = False
            rems = rems_grupos.get(c['Codigo_Interno'], [])
            for r in rems:
                f_r = pd.to_datetime(r['Fecha_Remedicion'])
                if f_r.month == m_idx and f_r.year == a and r.get('Baja_Pasivo', 0.0) > 0:
                    paso_terminacion_parcial = True
                    break

            # Asiento Automático por Término Natural del Contrato
            f_fin_c = pd.to_datetime(c['Fin'])
            if not paso_baja_manual and not paso_terminacion_parcial and c['Estado'] == 'Activo':
                if f_fin_c.month == m_idx and f_fin_c.year == a:
                    pasado = tab[tab['Fecha'] <= f_fin_c]
                    if not pasado.empty:
                        tc_baja = obtener_tc_cache(c['Moneda'], f_fin_c)
                        s_fin_pasivo = pasado.iloc[-1]['S_Fin_Orig'] * tc_baja
                        amort_acum = pasado['Dep_Orig'].sum() * tc_baja
                        s_fin_rou = (rou * tc_baja) - amort_acum
                        
                        if s_fin_pasivo > 0.01 or abs(rou * tc_baja) > 0.01:
                            t_baja = "6. Baja por Término Natural del Contrato"
                            if s_fin_pasivo > 0:
                                add_asiento(detalles, c['Empresa'], c['Codigo_Interno'], t_baja, *cta_map['Pasivo'], s_fin_pasivo, 0)
                            
                            r_neto = (rou * tc_baja)
                            add_asiento(detalles, c['Empresa'], c['Codigo_Interno'], t_baja, *cta_map['ROU'], 0, r_neto)
                            add_asiento(detalles, c['Empresa'], c['Codigo_Interno'], t_baja, *cta_map['AmortAcum'], amort_acum, 0)
                            
                            dif_baja = s_fin_pasivo - s_fin_rou
                            if dif_baja > 0:
                                add_asiento(detalles, c['Empresa'], c['Codigo_Interno'], t_baja, *cta_map['Ganancia'], 0, dif_baja)
                            elif dif_baja < 0:
                                add_asiento(detalles, c['Empresa'], c['Codigo_Interno'], t_baja, *cta_map['Perdida'], abs(dif_baja), 0)

            # Asiento de Ajuste por Remedición
            rems = rems_grupos.get(c['Codigo_Interno'], [])
            for r in rems:
                f_r = pd.to_datetime(r['Fecha_Remedicion'])
                if f_r.month == m_idx and f_r.year == a:
                    t_rem = "7. Ajuste por Remedición"
                    past_tab = tab[tab['Fecha'] < f_r]
                    fut_tab = tab[tab['Fecha'] >= f_r]
                    old_pasivo = past_tab.iloc[-1]['S_Fin_Orig'] if not past_tab.empty else vp
                    new_pasivo = fut_tab.iloc[0]['S_Ini_Orig'] if not fut_tab.empty else 0.0
                    
                    tc_rem = obtener_tc_cache(c['Moneda'], f_r)
                    if tc_rem == 0: tc_rem = 1.0
                    
                    # 1) Reconocer Terminación Parcial (Reducción de Alcance) si aplica
                    baja_p = r.get('Baja_Pasivo', 0.0)
                    baja_r = r.get('Baja_ROU', 0.0)
                    
                    if baja_p > 0 or baja_r > 0:
                        t_term = "7.a Terminación Parcial (Reducción de Alcance)"
                        bp_clp = baja_p * tc_rem
                        br_clp = baja_r * tc_ini
                        
                        # Calcular el diferencial matemático en CLP absoluto para cuadratura perfecta
                        pl_efecto_clp_real = bp_clp - br_clp
                        
                        add_asiento(detalles, c['Empresa'], c['Codigo_Interno'], t_term, *cta_map['Pasivo'], bp_clp, 0)
                        add_asiento(detalles, c['Empresa'], c['Codigo_Interno'], t_term, *cta_map['ROU'], 0, br_clp)
                        
                        if pl_efecto_clp_real > 0:
                            add_asiento(detalles, c['Empresa'], c['Codigo_Interno'], t_term, *cta_map['Ganancia'], 0, pl_efecto_clp_real)
                        elif pl_efecto_clp_real < 0:
                            add_asiento(detalles, c['Empresa'], c['Codigo_Interno'], t_term, *cta_map['Perdida'], abs(pl_efecto_clp_real), 0)
                            
                        old_pasivo -= baja_p
                    
                    # 2) Ajuste final ROU vs Pasivo
                    aj = (new_pasivo - old_pasivo) * tc_rem
                    
                    if abs(aj) > 0.01:
                        if aj > 0:
                            add_asiento(detalles, c['Empresa'], c['Codigo_Interno'], t_rem, *cta_map['ROU'], aj, 0)
                            add_asiento(detalles, c['Empresa'], c['Codigo_Interno'], t_rem, *cta_map['Pasivo'], 0, aj)
                        elif aj < 0:
                            add_asiento(detalles, c['Empresa'], c['Codigo_Interno'], t_rem, *cta_map['Pasivo'], abs(aj), 0)
                            add_asiento(detalles, c['Empresa'], c['Codigo_Interno'], t_rem, *cta_map['ROU'], 0, abs(aj))

        st.session_state.asientos_data = detalles
        st.session_state.asientos_params = {'m': m_nom, 'a': a}

    if 'asientos_data' in st.session_state:
        detalles = st.session_state.asientos_data
        m_saved = st.session_state.asientos_params['m']
        a_saved = st.session_state.asientos_params['a']
        
        t1, t2 = st.tabs(["Resumen Mensual Contable", "Detalle por Contrato"])
        
        with t2:
            if detalles:
                df_asientos = pd.DataFrame(detalles)
                df_asientos.rename(columns={'Cod1': 'ID_Contrato'}, inplace=True)
                # Agregar Fila Total al Detalle
                df_asientos.loc['Total'] = df_asientos.sum(numeric_only=True)
                df_asientos.at['Total', 'Empresa'] = 'TOTALES'
                df_asientos.at['Total', 'Cuenta'] = 'CUADRATURA'
                
                st.dataframe(df_asientos.style.format(precision=0, thousands="."))
                st.download_button("Exportar Detalle de Asientos (Excel)", to_excel(df_asientos), f"Detalle_Asientos_{m_saved}_{a_saved}.xlsx")
            else:
                st.info("No hay movimientos granulares para esta selección.")
                
        with t1:
            if detalles:
                df_asientos = pd.DataFrame(detalles)
                df_resumen = df_asientos.groupby(['Empresa', 'Transacción', 'N° Cuenta', 'Cuenta', 'Tipo']).sum(numeric_only=True).reset_index()
                # Fila de cuadratura
                df_resumen.loc['Total'] = df_resumen.sum(numeric_only=True)
                df_resumen.at['Total', 'Empresa'] = 'TOTALES'
                df_resumen.at['Total', 'Transacción'] = ''
                df_resumen.at['Total', 'N° Cuenta'] = ''
                df_resumen.at['Total', 'Cuenta'] = 'CUADRATURA PERFECTA'
                df_resumen.at['Total', 'Tipo'] = ''
                st.dataframe(df_resumen.style.format(precision=0, thousands="."))
                st.download_button("Exportar Asientos Resumidos (Excel)", to_excel(df_resumen), f"Resumen_Asientos_{m_saved}_{a_saved}.xlsx")
                
                from db import obtener_erp_activo, leer_credencial_erp
                erp_act = obtener_erp_activo()
                if erp_act:
                    st.markdown("---")
                    if st.button(f"🚀 Contabilizar Directamente en {erp_act}", type="primary", use_container_width=True):
                        st.info(f"🔄 Iniciando protocolo de conexión con {erp_act}...")
                        try:
                            # ---------------------------------------------------------
                            # THE LAST MILE (ESPACIO RESERVADO PARA INGENIERÍA IT)
                            # ---------------------------------------------------------
                            creds = leer_credencial_erp(erp_act)['secretos']
                            
                            if erp_act == "Odoo":
                                import xmlrpc.client
                                url, db, user, pwd = creds['url'], creds['db'], creds['user'], creds['pass']
                                
                                # Filtrar fila de totales
                                df_send = df_resumen[df_resumen['Empresa'] != 'TOTALES']
                                
                                line_ids = []
                                for _, row in df_send.iterrows():
                                    cuenta = str(row['N° Cuenta']).strip()
                                    debe = float(row['Debe']) if 'Debe' in row else 0.0
                                    haber = float(row['Haber']) if 'Haber' in row else 0.0
                                    nombre_cta = str(row['Cuenta']).strip()
                                    
                                    if debe == 0.0 and haber == 0.0:
                                        continue
                                        
                                    line_ids.append((0, 0, {
                                        'name': f"{row['Transacción']} - {nombre_cta}",
                                        'debit': debe,
                                        'credit': haber,
                                        '_codigo_cta': cuenta
                                    }))
                                
                                try:
                                    common = xmlrpc.client.ServerProxy(f'{url.rstrip("/")}/xmlrpc/2/common')
                                    uid = common.authenticate(db, user, pwd, {})
                                    if not uid:
                                        st.error("❌ Falló la autenticación con Odoo. Revise las credenciales en 'Configuración'.")
                                    else:
                                        models = xmlrpc.client.ServerProxy(f'{url.rstrip("/")}/xmlrpc/2/object')
                                        
                                        # Buscar el Diario por defecto (Misceláneos)
                                        j_ids = models.execute_kw(db, uid, pwd, 'account.journal', 'search', [[['type', '=', 'general']]], {'limit': 1})
                                        journal_id = j_ids[0] if j_ids else False
                                        
                                        # Resolver IDs de cuentas según el código para evitar errores relacionales
                                        codigos = list(set([l[2]['_codigo_cta'] for l in line_ids]))
                                        cuentas_odoo = models.execute_kw(db, uid, pwd, 'account.account', 'search_read', 
                                            [[['code', 'in', codigos]]], {'fields': ['id', 'code']})
                                        mapa_cuentas = {c['code']: c['id'] for c in cuentas_odoo}
                                        
                                        falta_cuenta = False
                                        for l in line_ids:
                                            cod = l[2].pop('_codigo_cta')
                                            acc_id = mapa_cuentas.get(cod)
                                            if acc_id:
                                                l[2]['account_id'] = acc_id
                                            else:
                                                st.warning(f"⚠️ La cuenta '{cod}' no existe en el plan de cuentas de Odoo.")
                                                falta_cuenta = True
                                        
                                        if not falta_cuenta:
                                            # Payload maestro del asiento
                                            move_vals = {
                                                'ref': f"IFRS 16 - {m_saved} {a_saved}",
                                                'date': pd.to_datetime(f"{a_saved}-{MESES_LISTA.index(m_saved)+1}-01") + relativedelta(day=31),
                                                'journal_id': journal_id,
                                                'line_ids': line_ids
                                            }
                                            
                                            # Formatear fecha para el JSON de Odoo (string)
                                            move_vals['date'] = move_vals['date'].strftime('%Y-%m-%d')
                                            
                                            # Crear Asiento Borrador
                                            move_id = models.execute_kw(db, uid, pwd, 'account.move', 'create', [move_vals])
                                            st.success(f"✅ ¡Éxito! Asiento preliminar (Borrador) creado en Odoo. (ID Asiento: {move_id})")
                                            
                                except Exception as e:
                                    st.error(f"❌ Ocurrió un error al enviar información a Odoo: {str(e)}")
                                
                            elif erp_act == "SAP ERP (OData/BAPI)":
                                # import requests
                                # response = requests.post(creds['url'] + '/API_JOURNALENTRYITEM', json=df_resumen.to_dict(), auth=(creds['user'], creds['pass']))
                                st.warning("⚠️ Pendiente IT: Inserte la mutación POST hacia SAP OData/Service Layer.")
                                
                            elif erp_act == "Microsoft Dynamics 365":
                                # import msal, requests
                                # app = msal.ConfidentialClientApplication(creds['client'], authority=f"https://login.microsoftonline.com/{creds['tenant']}", client_credential=creds['secret'])
                                # result = app.acquire_token_for_client(scopes=[f"{creds['env']}/.default"])
                                # requests.post(f"{creds['env']}/data/JournalEntries", headers={'Authorization': 'Bearer ' + result['access_token']}, json=...)
                                st.warning("⚠️ Pendiente IT: Inserte protocolo MSAL OAuth 2.0 y POST de Dynamics 365.")
                                
                            elif erp_act == "Oracle NetSuite":
                                # from requests_oauthlib import OAuth1
                                # auth = OAuth1(creds['ckey'], creds['csec'], creds['tkey'], creds['tsec'], signature_method='HMAC-SHA256', realm=creds['account'])
                                # requests.post("https://{}.suitetalk.api.netsuite.com/services/rest/record/v1/journalentry".format(creds['account']), auth=auth, json=...)
                                st.warning("⚠️ Pendiente IT: Inserte Cliente SuiteTalk REST/SOAP para NetSuite usando OAuth1.")
                                
                            st.success(f"Simulacro exitoso: El conector detectó {len(df_resumen)-1} líneas de asiento preparadas para enviar.")
                        except Exception as e:
                            st.error(f"Error forzado de conexión: {e}")
            else:
                st.info("No hay movimientos para esta selección.")

def modulo_notas():
    st.header("📋 Movimiento de saldos (Nota de Pasivos y Activos)")
    c1, c2, c3 = st.columns(3)
    emp_sel = c1.selectbox("Empresa", ["Todas"] + EMPRESAS_LISTA, key="nt_emp")
    m_nom = c2.selectbox("Mes Nota", MESES_LISTA, key="nt_m")
    a = c3.number_input("Año Nota", value=date.today().year, key="nt_a")
    
    if st.button("Generar Movimiento de saldos", type="primary"):
        m_idx = MESES_LISTA.index(m_nom) + 1
        f_act = pd.to_datetime(date(a, m_idx, 1)) + relativedelta(day=31)
        f_ant = pd.to_datetime(date(a - 1, 12, 31))
        
        roll_pasivo = []
        roll_activo = []
        
        lista_c = cargar_contratos()
        from db import cargar_remediciones_todas_agrupadas
        rems_grupos = cargar_remediciones_todas_agrupadas()
        
        if emp_sel != "Todas": lista_c = [c for c in lista_c if c['Empresa'] == emp_sel]
            
        for c in lista_c:
            f_ini_c = pd.to_datetime(c['Inicio'])
            if f_act < f_ini_c.replace(day=1): continue
            
            tab, vp, rou = obtener_motor_financiero(c, rems=rems_grupos.get(c['Codigo_Interno'], []))
            if tab.empty or 'Fecha' not in tab.columns: continue
            tc_act = obtener_tc_cache(c['Moneda'], f_act)
            tc_ant = obtener_tc_cache(c['Moneda'], f_ant)
            r_act_pasivo = tc_act 
            r_ant_pasivo = tc_ant 
            
            tc_ini = float(c['Valor_Moneda_Inicio']) if c.get('Valor_Moneda_Inicio') and float(c['Valor_Moneda_Inicio']) > 0 else 1.0
            r_act_rou = tc_ini
            r_ant_rou = tc_ini
            
            f_ini_c = pd.to_datetime(c['Inicio'])
            fue_adicionado = (f_ini_c.year == a)
            
            # Saldo Inicial (Cierre del año anterior YTD)
            past_ant = tab[tab['Fecha'] <= f_ant]
            if fue_adicionado:
                s_ini = 0
                s_ini_rou = 0
                s_ini_orig = 0
                s_ini_rou_orig = 0
            else:
                s_ini_orig = past_ant.iloc[-1]['S_Fin_Orig'] if not past_ant.empty else 0
                s_ini = s_ini_orig * r_ant_pasivo
                
                s_ini_rou_orig = rou - (past_ant['Dep_Orig'].sum() if not past_ant.empty else 0)
                s_ini_rou = s_ini_rou_orig * r_ant_rou
            
            # Reconocimiento Inicial (Adiciones YTD)
            if fue_adicionado:
                adic_pasivo = vp * tc_ini
                adic_rou = rou * tc_ini
            else:
                adic_pasivo = 0
                adic_rou = 0
                
            # Movimientos del año (YTD) hasta el mes seleccionado
            curr = tab[(tab['Fecha'].dt.year == a) & (tab['Fecha'].dt.month <= m_idx)]
            interes = curr['Int_Orig'].sum() * r_act_pasivo if not curr.empty else 0
            pagos = curr['Pago_Orig'].sum() * r_act_pasivo if not curr.empty else 0
            amortizacion = curr['Dep_Orig'].sum() * r_act_rou if not curr.empty else 0
            
            # 4. Modificadores (Bajas, Remediciones) 
            rems = rems_grupos.get(c['Codigo_Interno'], [])
            bajas_p, bajas_a = 0, 0
            rem_p, rem_a = 0, 0
            
            # Las remediciones sumarán el valor del ajuste ROU inyectado en ese periodo
            for r in rems:
                f_r = pd.to_datetime(r['Fecha_Remedicion'])
                # Filtramos que haya ocurrido en el YTD hasta el mes evaluado
                if f_r.year == a and f_r.month <= m_idx:
                    past_tab_rem = tab[tab['Fecha'] < f_r]
                    fut_tab_rem = tab[tab['Fecha'] >= f_r]
                    old_pasivo_rem = past_tab_rem.iloc[-1]['S_Fin_Orig'] if not past_tab_rem.empty else vp
                    new_pasivo_rem = fut_tab_rem.iloc[0]['S_Ini_Orig'] if not fut_tab_rem.empty else 0.0
                    
                    tc_rem_spot = obtener_tc_cache(c['Moneda'], f_r)
                    if tc_rem_spot == 0: tc_rem_spot = 1.0
                    
                    salto_rem = (new_pasivo_rem - old_pasivo_rem) * tc_rem_spot
                    rem_a += salto_rem
                    rem_p += salto_rem
                    
            # 5. Fijación Excéntrica del Saldo Final (Fotografía Real)
            curr_fin = tab[tab['Fecha'] <= f_act]
            if not curr_fin.empty:
                s_fin_orig_real = curr_fin.iloc[-1]['S_Fin_Orig']
                dep_acum_real = curr_fin['Dep_Orig'].sum()
            else:
                s_fin_orig_real = vp
                dep_acum_real = 0
                
            s_fin_real = s_fin_orig_real * r_act_pasivo
            s_fin_rou_real = (rou * r_act_rou) + rem_a - (dep_acum_real * r_act_rou)

            # Validar Bajas Definitivas consolidadas y vencimientos naturales en YTD
            es_baja = False
            f_fin_c = pd.to_datetime(c['Fin'])
            f_baja_real = None

            if c.get('Fecha_Baja') and c['Estado'] == 'Baja':
                f_baja_real = pd.to_datetime(c['Fecha_Baja'])
            elif f_act.year > f_fin_c.year or (f_act.year == f_fin_c.year and f_act.month >= f_fin_c.month):
                f_baja_real = f_fin_c

            if f_baja_real:
                if f_baja_real.year < a:
                    continue # Expiró el año anterior, excluir completamente
                elif f_baja_real.year == a and f_baja_real.month <= m_idx:
                    es_baja = True
                    bajas_p = -s_fin_real
                    bajas_a = -s_fin_rou_real
                    s_fin_real = 0
                    s_fin_rou_real = 0
                    
            # 6. Recalcular la Diferencia de Cambio como "Plug" (Cuadratura Perfecta)
            # Puente Teórico = S_Ini + Adiciones + Interes - Pagos + Rem_P + Dif_Cambio = S_Fin
            # Despejando -> Dif_Cambio = S_Fin - S_Ini - Adiciones - Interes + Pagos - Rem_P - Bajas
            reajuste = s_fin_real - s_ini - adic_pasivo - interes + pagos - rem_p - bajas_p
            reajuste_rou = s_fin_rou_real - s_ini_rou - adic_rou - rem_a + amortizacion - bajas_a
            
            if rems:
                last_rem = rems[-1]
                n_can, n_tas, n_plaz, n_fin = last_rem['Canon'], last_rem['Tasa']*100, last_rem['Plazo'], last_rem['Fin']
            else:
                n_can, n_tas, n_plaz, n_fin = None, None, None, None
                
            fin_orig = (pd.to_datetime(c['Inicio']) + pd.DateOffset(months=int(c['Plazo']))).strftime('%Y-%m-%d')
                    
            roll_pasivo.append({"ID_Contrato": c['Codigo_Interno'], "Empresa": c['Empresa'], "Clase_Activo": c['Clase_Activo'], "Contrato": c['Nombre'], "S.Inicial": s_ini, "Adiciones": adic_pasivo, "Remediciones": rem_p, "Nuevo_Canon": n_can, "Nueva_Tasa_Anual_%": n_tas, "Nuevo_Plazo": n_plaz, "Nuevo_Fin": n_fin, "Fin_Original": fin_orig, "Interés devengado": interes, "Dif. Cambio": reajuste, "Pagos de capital": -(pagos - interes), "Pago de intereses": -interes, "Bajas": bajas_p, "S.Final": s_fin_real})
            roll_activo.append({"ID_Contrato": c['Codigo_Interno'], "Empresa": c['Empresa'], "Clase_Activo": c['Clase_Activo'], "Contrato": c['Nombre'], "S.Inicial": s_ini_rou, "Adiciones": adic_rou, "Remediciones": rem_a, "Nuevo_Canon": n_can, "Nueva_Tasa_Anual_%": n_tas, "Nuevo_Plazo": n_plaz, "Nuevo_Fin": n_fin, "Fin_Original": fin_orig, "Amortización": -amortizacion, "Bajas": bajas_a, "S.Final": s_fin_rou_real})
        
        st.session_state.roll_pasivo = roll_pasivo
        st.session_state.roll_activo = roll_activo
        st.session_state.roll_params = {'m': m_nom, 'a': a}

    if 'roll_pasivo' in st.session_state and 'roll_activo' in st.session_state:
        roll_pasivo = st.session_state.roll_pasivo
        roll_activo = st.session_state.roll_activo
        m_saved = st.session_state.roll_params['m']
        a_saved = st.session_state.roll_params['a']
        
        t1, t2 = st.tabs(["Movimiento de saldos por Clase", "Detalle por Contrato individual"])
        
        with t1:
            st.subheader("Movimiento de saldos (Horizontal) - Pasivos")
            if roll_pasivo:
                df_pas = pd.DataFrame(roll_pasivo)
                cols_p_sum = ['S.Inicial', 'Adiciones', 'Remediciones', 'Interés devengado', 'Dif. Cambio', 'Pagos de capital', 'Pago de intereses', 'Bajas', 'S.Final']
                res_pasivo = df_pas.groupby('Clase_Activo')[cols_p_sum].sum()
                res_pasivo.loc['TOTAL PASIVO'] = res_pasivo.sum()
                
                # Apply smaller font size style if necessary
                styled_pas = res_pasivo.style.format(precision=0, thousands=".").set_properties(**{'font-size': '14px', 'white-space': 'nowrap'})
                st.dataframe(styled_pas, use_container_width=True)
                
                st.subheader("Movimiento de saldos (Horizontal) - Activos ROU")
                df_act = pd.DataFrame(roll_activo)
                cols_a_sum = ['S.Inicial', 'Adiciones', 'Remediciones', 'Amortización', 'Bajas', 'S.Final']
                res_activo = df_act.groupby('Clase_Activo')[cols_a_sum].sum()
                res_activo.loc['TOTAL ACTIVO'] = res_activo.sum()
                
                styled_act = res_activo.style.format(precision=0, thousands=".").set_properties(**{'font-size': '14px', 'white-space': 'nowrap'})
                st.dataframe(styled_act, use_container_width=True)
                
            else:
                st.info("No hay datos")
    
        with t2:
            if roll_pasivo:
                st.subheader("Pasivos por Arrendamiento (Detalle)")
                st.dataframe(df_pas.style.format(precision=0, thousands="."))
                st.download_button("Exportar Pasivos", to_excel(df_pas), f"RPasivos_{m_saved}_{a_saved}.xlsx")
                
                st.subheader("Activos por Derecho de Uso ROU (Detalle)")
                st.dataframe(df_act.style.format(precision=0, thousands="."))
                st.download_button("Exportar Activos", to_excel(df_act), f"RActivos_{m_saved}_{a_saved}.xlsx")

def modulo_dashboard():
    st.header("🧮 Panel de Saldos")
    
    c1, c2, c3 = st.columns(3)
    empresas = ["Todas"] + EMPRESAS_LISTA
    emp_sel = c1.selectbox("Empresa", empresas, key="dash_emp")
    m = c2.selectbox("Mes", MESES_LISTA, key="dash_m")
    a = c3.number_input("Año", value=date.today().year, key="dash_a")
    
    if st.button("Generar Resumen de Saldos", type="primary"):
        f_t = pd.to_datetime(date(a, MESES_LISTA.index(m)+1, 1)) + relativedelta(day=31)
        df_c = pd.DataFrame(cargar_contratos())
        from db import cargar_remediciones_todas_agrupadas
        rems_grupos = cargar_remediciones_todas_agrupadas()
        
        if df_c.empty:
            st.session_state.dash_data = None
            st.warning("No hay contratos registrados.")
        else:
            res = []
            for _, c in df_c.iterrows():
                if emp_sel != "Todas" and c['Empresa'] != emp_sel: continue
                if f_t < pd.to_datetime(c['Inicio']).replace(day=1): continue
                
                es_baja_ejercicio = False
                f_baja_efectiva = None
                
                # 1. Identificar si existe baja anticipada
                if c.get('Fecha_Baja') and c['Estado'] == 'Baja':
                    f_baja = pd.to_datetime(c['Fecha_Baja'])
                    if f_baja <= f_t:
                        f_baja_efectiva = f_baja
                
                # 2. Identificar si expiró naturalmente
                f_fin_c = pd.to_datetime(c['Fin'])
                if f_t.year > f_fin_c.year or (f_t.year == f_fin_c.year and f_t.month >= f_fin_c.month):
                    if not f_baja_efectiva or f_fin_c < f_baja_efectiva:
                        f_baja_efectiva = f_fin_c
                        
                if f_baja_efectiva:
                    if f_baja_efectiva.year < a:
                        continue # Excluir categóricamente si murió antes del año de consulta
                    elif f_baja_efectiva.year == a and f_baja_efectiva.month <= f_t.month:
                        es_baja_ejercicio = True # Incluir en Detalle con balances a 0 para preservar P&L
                
                tab, vp, rou = obtener_motor_financiero(c, rems=rems_grupos.get(c['Codigo_Interno'], []))
                if tab.empty or 'Fecha' not in tab.columns: continue
                past = tab[tab['Fecha'] <= f_t]
                if not past.empty:
                    tc = obtener_tc_cache(c['Moneda'], f_t); ratio_pasivo = tc
                    v_act = past.iloc[-1]['S_Fin_Orig']
                    
                    futuros = tab[tab['Fecha'] > f_t].copy()
                    v_cor_sum = 0
                    v_noc_sum_raw = 0
                    if not futuros.empty:
                        limite_12_dash = f_t + relativedelta(months=12)
                        
                        # Capital puro del periodo = S_Ini_Orig - S_Fin_Orig
                        futuros['Capital'] = futuros['S_Ini_Orig'] - futuros['S_Fin_Orig']
                        
                        # La ultima fila debe agregar todo el S_Fin_Orig remanente a su capital (ajuste del motor original)
                        futuros.iloc[-1, futuros.columns.get_loc('Capital')] += futuros.iloc[-1]['S_Fin_Orig']
                        
                        dias_al_pago = (futuros['Fecha'] - f_t).dt.days
                        es_corriente_mask = (dias_al_pago <= 90) | (futuros['Fecha'] <= limite_12_dash)
                        
                        v_cor_sum = futuros.loc[es_corriente_mask, 'Capital'].sum()
                        v_noc_sum_raw = futuros.loc[~es_corriente_mask, 'Capital'].sum()
                                
                    # El Pasivo Total debe ser exactamente el balance de cierre actual (v_act)
                    # El Pasivo Corriente son las amortizaciones estrictas de los proximos 12 meses
                    # El Pasivo No Corriente absorbe cualquier diferencia (ej. remediciones futuras)
                    v12 = v_act - v_cor_sum 
                    
                    tc_ini = float(c['Valor_Moneda_Inicio']) if float(c['Valor_Moneda_Inicio']) > 0 else 1.0
                    rou_bruto = rou * tc_ini
                    
                    # Using the preloaded rems_grupos cache to prevent N+1 queries
                    rems = rems_grupos.get(c['Codigo_Interno'], [])
                    
                    n_can, n_tas, n_plaz, n_fin, n_rou = None, None, None, None, None
                    if rems:
                         last_rem = rems[-1]
                         n_can = last_rem['Canon']
                         n_tas = last_rem['Tasa'] * 100
                         n_plaz = last_rem['Plazo']
                         n_fin = last_rem['Fin']
                         
                         f_r_last = pd.to_datetime(last_rem['Fecha_Remedicion'])
                         fut_r_last = tab[tab['Fecha'] >= f_r_last]
                         n_rou_orig = fut_r_last.iloc[0]['S_Ini_Orig'] if not fut_r_last.empty else 0.0
                         tc_rem_last = obtener_tc_cache(c['Moneda'], f_r_last)
                         if tc_rem_last == 0: tc_rem_last = 1.0
                         n_rou = n_rou_orig * tc_rem_last
                         
                    for r in rems:
                        f_r = pd.to_datetime(r['Fecha_Remedicion'])
                        if f_r <= f_t:
                            tc_rem = obtener_tc_cache(c['Moneda'], f_r)
                            if tc_rem == 0: tc_rem = 1.0
                            
                            past_r = tab[tab['Fecha'] < f_r]
                            fut_r = tab[tab['Fecha'] >= f_r]
                            old_pasivo = past_r.iloc[-1]['S_Fin_Orig'] if not past_r.empty else vp
                            new_pasivo = fut_r.iloc[0]['S_Ini_Orig'] if not fut_r.empty else 0.0
                            
                            baja_p = r.get('Baja_Pasivo', 0.0)
                            baja_r = r.get('Baja_ROU', 0.0)
                            
                            old_pasivo_adjusted = old_pasivo - baja_p
                            jump_rou_uf = new_pasivo - old_pasivo_adjusted
                            
                            if baja_r > 0:
                                rou_bruto -= (baja_r * tc_ini)
                            
                            if abs(jump_rou_uf) > 0.01:
                                rou_bruto += (jump_rou_uf * tc_rem)

                    amort_acum = past['Dep_Orig'].sum()
                    amort_clp = amort_acum * tc_ini

                    past_ejercicio = past[past['Fecha'].dt.year == a]
                    dep_ejercicio_clp = past_ejercicio['Dep_Orig'].sum() * tc_ini
                    
                    if es_baja_ejercicio:
                        v_act = 0
                        v_cor_sum = 0
                        v12 = 0
                        rou_bruto = 0
                        amort_clp = 0
                        rou = 0

                    item_dict = {}
                    
                    # 1. Estado de Vigencia al Corte
                    f_fin_date = pd.to_datetime(c['Fin']).date()
                    f_t_date = f_t.date()
                    if c.get('Estado') == 'Baja' and c.get('Fecha_Baja'):
                        estado_vig = '🚨 Dado de Baja'
                        estado_real = 'Dado de Baja'
                    elif f_fin_date <= f_t_date:
                        estado_vig = '🚨 Expirado'
                        estado_real = 'Expirado'
                    else:
                        estado_vig = '🟢 Vigente'
                        estado_real = 'Activo'
                        
                    item_dict["Estado Vigencia al Corte"] = estado_vig
                    
                    # 2. Copia exacta de columnas oficiales y extras activos (evitar campos zombie borrados)
                    columnas_base = ['Codigo_Interno', 'Empresa', 'Clase_Activo', 'ID', 'Proveedor', 'Nombre', 'Moneda', 'Canon', 'Tasa', 'Tasa_Mensual', 'Valor_Moneda_Inicio', 'Plazo', 'Inicio', 'Fin', 'Estado', 'Fecha_Baja', 'Ajuste_ROU', 'Tipo_Pago', 'Fecha_Remedicion', 'Frecuencia_Pago']
                    permitidas = columnas_base + obtener_parametros('CAMPO_EXTRA')
                    
                    for k, v in c.items():
                        if k in permitidas:
                            if k == 'Tasa':
                                item_dict['Tasa Anual %'] = v
                            elif k == 'Estado':
                                item_dict[k] = estado_real
                            else:
                                item_dict[k] = v
                                
                            if k == 'Fin':
                                item_dict["Fin_Original"] = (pd.to_datetime(c['Inicio']) + pd.DateOffset(months=int(c['Plazo']))).strftime('%Y-%m-%d')
                                
                    item_dict["Nuevo_Canon"] = n_can
                    item_dict["Nueva_Tasa_Anual_%"] = n_tas
                    item_dict["Nuevo_Plazo"] = n_plaz
                    item_dict["Nuevo_Fin"] = n_fin
                    item_dict["Nuevo_Activo_ROU"] = n_rou
                        
                    # 3. Columnas de cálculos financieros del Dashboard (se agregan al final)
                    item_dict["Cuotas Devengadas"] = len(past)
                    item_dict["Cuotas por Devengar"] = len(tab) - len(past)
                    item_dict["Cuotas de Pago (VA)"] = len(tab[tab['Pago_Orig'] > 0])
                    item_dict["Valor Inicial ROU"] = rou * tc_ini
                    item_dict["Depreciación Ejercicio"] = dep_ejercicio_clp
                    item_dict["ROU Bruto"] = rou_bruto
                    item_dict["Amort. Acum"] = amort_clp
                    item_dict["ROU Neto"] = rou_bruto - amort_clp
                    item_dict["Pasivo Total"] = v_act * ratio_pasivo
                    item_dict["Pasivo Corriente"] = v_cor_sum * ratio_pasivo
                    item_dict["Pasivo No Corr"] = v12 * ratio_pasivo
                        
                    res.append(item_dict)
            
            if res:
                st.session_state.dash_data = pd.DataFrame(res)
                st.session_state.dash_params = {'m': m, 'a': a, 'emp': emp_sel}
            else:
                st.session_state.dash_data = pd.DataFrame() # Vacío pero instanciado
                st.session_state.dash_params = {'m': m, 'a': a, 'emp': emp_sel}
    
    # Renderizado condicional basado en session_state
    if 'dash_data' in st.session_state and st.session_state.dash_data is not None:
        df_res = st.session_state.dash_data
        m_saved = st.session_state.dash_params['m']
        a_saved = st.session_state.dash_params['a']
        
        t1, t2 = st.tabs(["Resumen", "Detalle por Contrato"])
        
        with t1:
            if not df_res.empty:
                cols_to_sum = ['Valor Inicial ROU', 'ROU Bruto', 'Amort. Acum', 'ROU Neto', 'Pasivo Total', 'Pasivo Corriente', 'Pasivo No Corr']
                df_grp = df_res.groupby('Empresa')[cols_to_sum].sum(numeric_only=True).reset_index()
                
                # Agregar fila de TOTALES si están "Todas" seleccionadas (es decir, hay más de 1 empresa o se seleccionó 'Todas')
                if len(df_grp) > 1 or st.session_state.dash_params.get('emp') == "Todas":
                    total_row = df_grp[cols_to_sum].sum().to_frame().T
                    total_row['Empresa'] = 'TOTALES'
                    df_grp = pd.concat([df_grp, total_row], ignore_index=True)

                # Eliminar "Valor Inicial ROU" de la vista consolidada a pedido del usuario
                if 'Valor Inicial ROU' in df_grp.columns:
                    df_grp = df_grp.drop(columns=['Valor Inicial ROU'])

                st.dataframe(df_grp.style.format(precision=0, thousands="."))
                st.download_button("Exportar Resumen (Excel)", to_excel(df_grp), f"Resumen_Saldos_{m_saved}_{a_saved}.xlsx")
            else:
                st.info("No hay datos para esta selección.")
                
        with t2:
            if not df_res.empty:
                numeric_cols = df_res.select_dtypes(include='number').columns
                cols_clp = [c for c in numeric_cols if c in ['Valor Inicial ROU', 'Depreciación Ejercicio', 'ROU Bruto', 'Amort. Acum', 'ROU Neto', 'Pasivo Total', 'Pasivo Corriente', 'Pasivo No Corr']]
                cols_canon = [c for c in numeric_cols if c in ['Canon', 'Nuevo_Canon', 'Valor_Moneda_Inicio']]
                
                # Format CLP fields as integers (precision=0)
                styled_df = df_res.style.format(precision=0, thousands=".", subset=cols_clp)
                # Format Canon and UF with decimals
                styled_df = styled_df.format(precision=4, thousands=".", subset=cols_canon)
                st.dataframe(styled_df)
                st.download_button("Exportar Detalle (Excel)", to_excel(df_res), f"Resumen_Saldos_Detalle_{m_saved}_{a_saved}.xlsx")

def modulo_monedas():
    st.header("💱 Monedas")
    st.warning("⚠️ Importante: Los valores de las monedas no pueden exceder los 4 decimales.")
    
    if "Lector" in st.session_state.rol:
        st.info("Rol Lector: Solo puedes visualizar los tipos de cambio.")
        t_unica = st.tabs(["Ver Todos los Datos"])[0]
        with t_unica:
            st.subheader("Histórico Completo de Tipos de Cambio")
            st.dataframe(cargar_monedas(), use_container_width=True)
        return
        
    t1, t2, t3 = st.tabs(["Carga Manual", "Carga Masiva", "Ver Todos los Datos"])
    
    with t1:
        monedas_activas = obtener_parametros('MONEDA')
        if not monedas_activas: monedas_activas = ["UF", "CLP", "USD", "EUR"]
        f, m, v = st.date_input("Fecha", min_value=date(1900, 1, 1), max_value=date(2100, 12, 31)), st.selectbox("Moneda", monedas_activas), st.number_input("Valor Moneda", format="%.4f")
        if st.button("Guardar Moneda"): 
            insertar_moneda(f.strftime('%Y-%m-%d'), m, v)
            st.cache_data.clear()
            if 'motor_cache' in st.session_state: st.session_state.motor_cache.clear()
            st.session_state.success_msg = f"Moneda {m} guardada exitosamente para la fecha {f.strftime('%Y-%m-%d')}."
            st.rerun()
        
    with t2:
        st.subheader("Carga Masiva de Tipos de Cambio")
        plantilla_m = pd.DataFrame(columns=["fecha", "moneda", "valor"])
        st.download_button("Descargar Plantilla", to_excel(plantilla_m), "plantilla_monedas.xlsx")
        
        arch_m = st.file_uploader("Subir Plantilla de Monedas", type=["xlsx"])
        if arch_m is not None:
            if st.button("Cargar Tipos de Cambio"):
                try:
                    df_in_m = pd.read_excel(arch_m)
                    cargar_masivo_monedas(df_in_m)
                    st.cache_data.clear()
                    if 'motor_cache' in st.session_state: st.session_state.motor_cache.clear()
                    st.session_state.success_msg = f"Tipos de cambio cargados exitosamente. Se han cargado {len(df_in_m)} líneas."
                    st.rerun()
                except Exception as e:
                    st.error(f"Error procesando archivo de monedas: {e}")
                    
    with t3:
        st.subheader("Histórico Completo de Tipos de Cambio")
        st.dataframe(cargar_monedas(), use_container_width=True)

def modulo_contratos():
    st.header("📝 Contratos")
    
    if "Lector" in st.session_state.rol:
        st.info("Rol Lector: Solo puedes visualizar los contratos existentes.")
        t_unica = st.tabs(["Ver Todos los Datos"])[0]
        with t_unica:
            st.subheader("Base de Datos Completa: Contratos")
            lista_c = cargar_contratos()
            from db import cargar_remediciones_todas_agrupadas
            rems_grupos = cargar_remediciones_todas_agrupadas()
            for c in lista_c:
                c['Fin_Original'] = (pd.to_datetime(c['Inicio']) + pd.DateOffset(months=int(c['Plazo']))).strftime('%Y-%m-%d')
                rems = rems_grupos.get(c['Codigo_Interno'], [])
                if rems:
                    last_rem = rems[-1]
                    c['Nuevo_Canon'] = last_rem['Canon']
                    c['Nueva_Tasa_Anual_%'] = last_rem['Tasa'] * 100
                    c['Nuevo_Plazo'] = last_rem['Plazo']
                    c['Nuevo_Fin'] = last_rem['Fin']
                else:
                    c['Nuevo_Canon'] = None
                    c['Nueva_Tasa_Anual_%'] = None
                    c['Nuevo_Plazo'] = None
                    c['Nuevo_Fin'] = None
                    
            df_contratos = pd.DataFrame(lista_c)
            if not df_contratos.empty:
                ahoy = date.today()
                df_contratos.insert(1, "Estado Vigencia", df_contratos['Fin'].apply(
                    lambda x: '🚨 Vencido' if pd.to_datetime(x).date() < ahoy else '🟢 Vigente'
                ))
                columnas_base = ['Codigo_Interno', 'Empresa', 'Clase_Activo', 'ID', 'Proveedor', 'Nombre', 'Moneda', 'Canon', 'Tasa', 'Tasa_Mensual', 'Valor_Moneda_Inicio', 'Plazo', 'Inicio', 'Fin', 'Fin_Original', 'Estado', 'Fecha_Baja', 'Ajuste_ROU', 'Tipo_Pago', 'Fecha_Remedicion', 'Frecuencia_Pago']
                permitidas = ['Estado Vigencia'] + columnas_base + obtener_parametros('CAMPO_EXTRA') + ['Nuevo_Canon', 'Nueva_Tasa_Anual_%', 'Nuevo_Plazo', 'Nuevo_Fin']
                cols_to_keep = [col for col in permitidas if col in df_contratos.columns]
                
                df_display = df_contratos[cols_to_keep]
                if 'Tasa' in df_display.columns:
                    df_display = df_display.rename(columns={'Tasa': 'Tasa Anual %'})
                st.dataframe(df_display, use_container_width=True)
                st.download_button("Exportar Contratos (Excel)", to_excel(df_display), "Base_Contratos_Completa.xlsx")
            else:
                st.info("No hay contratos cargados.")
        return
        
    t1, t2, t3, t4, t5, t6 = st.tabs(["Ingreso de Contrato Manual", "Ingreso de Contrato Masiva", "Ver Todos los Datos", "Modificación Individual", "Baja Anticipada", "Modificación Masiva"])
    
    with t1:
        with st.form("f"):
            c1, c2, c3 = st.columns(3)
            emp = c1.selectbox("Empresa", EMPRESAS_LISTA)
            clase = c1.selectbox("Clase", CLASES_ACTIVO)
            prov = c1.text_input("Proveedor")
            nom, id_p = c1.text_input("Nombre Contrato"), c1.text_input("ID/RUT")
            
            monedas_activas = obtener_parametros('MONEDA')
            if not monedas_activas: monedas_activas = ["UF", "CLP", "USD", "EUR"]
            mon, can = c2.selectbox("Moneda", monedas_activas), c2.number_input("Canon", format="%.4f")
            tas = c2.number_input("Tasa Anual %", value=6.0)
            
            frecuencias_raw = obtener_parametros('FRECUENCIA_PAGO')
            if not frecuencias_raw: frecuencias_raw = ["Mensual|1", "Trimestral|3", "Semestral|6", "Anual|12"]
            nombres_frec = [f.split('-')[0] for f in frecuencias_raw]
            frec = c2.selectbox("Frecuencia Pago", nombres_frec)
            
            f_i, f_f = c3.date_input("Inicio", min_value=date(1900, 1, 1), max_value=date(2100, 12, 31)), c3.date_input("Fin", min_value=date(1900, 1, 1), max_value=date(2100, 12, 31))
            t_pago = c3.selectbox("Tipo de Pago", ["Vencido", "Anticipado"])
            
            # Cargar dinámicamente campos extra configurados
            campos_extra = obtener_parametros('CAMPO_EXTRA')
            extra_vals = {}
            if campos_extra:
                st.markdown("---")
                st.write("**Campos Adicionales**")
                cols_extra = st.columns(3)
                for idx, cx in enumerate(campos_extra):
                    extra_vals[cx] = cols_extra[idx % 3].text_input(cx, key=f"cx_man_{cx}")
            
            # Descripciones ocultadas a pedido del usuario
            cd = 0.0
            pa = 0.0
            cdesm = 0.0
            inc = 0.0
            
            if st.form_submit_button("Registrar"):
                diff = relativedelta(f_f, f_i)
                p = diff.years * 12 + diff.months
                if diff.days >= 15: p += 1
                nuevo_c = {
                    "Codigo_Interno": generar_codigo_correlativo(emp, cargar_contratos()), 
                    "Empresa": emp, "Clase_Activo": clase, "ID": id_p, "Proveedor": prov, 
                    "Cod1": "", "Cod2": "", "Nombre": nom, "Moneda": mon, "Canon": can, 
                    "Tasa": tas/100, "Tasa_Mensual": pow(1+tas/100, 1/12)-1, 
                    "Valor_Moneda_Inicio": obtener_tc_cache(mon, f_i), "Plazo": p, 
                    "Inicio": f_i.strftime('%Y-%m-%d'), "Fin": f_f.strftime('%Y-%m-%d'), 
                    "Estado": "Activo", "Tipo_Pago": t_pago, 
                    "Costos_Directos": cd, "Pagos_Anticipados": pa, 
                    "Costos_Desmantelamiento": cdesm, "Incentivos": inc,
                    "Frecuencia_Pago": frec
                }
                nuevo_c.update(extra_vals)
                insertar_contrato(nuevo_c)
                if 'motor_cache' in st.session_state: st.session_state.motor_cache.clear()
                st.session_state.success_msg = "Contrato creado manualmente"
                st.rerun()

    with t2:
        st.subheader("Carga Masiva de Contratos")
        st.write("❗ **Instrucciones para la plantilla:**")
        st.write("- **Formato de Fechas**: El formato de ingreso de las fechas (Inicio, Fin) debe ser SIEMPRE: YYYY-MM-DD (ej. 2026-03-01)")
        st.write("- **Empresa**: Debe coincidir exacto. Ej: Holdco, Pacifico")
        st.write("- **Clase_Activo**: Oficinas, Vehículos, Maquinaria, Bodegas, Inmuebles")
        monedas_activas = obtener_parametros('MONEDA')
        if not monedas_activas: monedas_activas = ["UF", "CLP", "USD", "EUR"]
        st.write(f"- **Moneda**: {', '.join(monedas_activas)}")
        st.write("- **Tipo_Pago**: Vencido, Anticipado")
        
        frecuencias_raw = obtener_parametros('FRECUENCIA_PAGO')
        if not frecuencias_raw: frecuencias_raw = ["Mensual|1", "Trimestral|3", "Semestral|6", "Anual|12"]
        nombres_frec = [f.split('-')[0] for f in frecuencias_raw]
        st.write(f"- **Frecuencia_Pago**: {', '.join(nombres_frec)}")
        
        campos_extra = obtener_parametros('CAMPO_EXTRA')
        if campos_extra:
            st.write(f"- **Campos Extra Obligatorios en el Excel**: {', '.join(campos_extra)}")
        
        # Plantilla con valores de ejemplo para guiar al usuario
        ej_1 = {"Empresa": "Holdco", "Clase_Activo": "Oficinas", "ID": "76.123.456-7", "Proveedor": "Inmobiliaria SPA", "Nombre": "Oficina Central", "Moneda": monedas_activas[0] if monedas_activas else "UF", "Canon": 150.5, "Tasa Anual %": 6.0, "Frecuencia_Pago": nombres_frec[0] if nombres_frec else "Mensual", "Tipo_Pago": "Vencido", "Inicio": "2026-01-01", "Fin": "2028-12-31", "Costos_Directos": 50.0, "Pagos_Anticipados": 0, "Costos_Desmantelamiento": 100, "Incentivos": 0}
        ej_2 = {"Empresa": "Pacifico", "Clase_Activo": "Vehículos", "ID": "Camioneta AB12", "Proveedor": "Automotriz Leasing S.A.", "Nombre": "Leasing Vehicular", "Moneda": monedas_activas[1] if len(monedas_activas)>1 else "CLP", "Canon": 500000, "Tasa Anual %": 4.5, "Frecuencia_Pago": nombres_frec[-1] if nombres_frec else "Mensual", "Tipo_Pago": "Anticipado", "Inicio": "2026-03-01", "Fin": "2029-03-01", "Costos_Directos": 0, "Pagos_Anticipados": 500000, "Costos_Desmantelamiento": 0, "Incentivos": 50000}
        
        for cx in campos_extra:
            ej_1[cx] = "Info Test 1"
            ej_2[cx] = "Info Test 2"
            
        df_ejemplo = pd.DataFrame([ej_1, ej_2])
        
        st.download_button("Descargar Plantilla con Ejemplos", to_excel(df_ejemplo), "plantilla_contratos.xlsx")
        
        arch = st.file_uploader("Subir Plantilla Completa", type=["xlsx"])
        if arch is not None:
            if st.button("Procesar y Cargar Masivamente"):
                try:
                    df_in = pd.read_excel(arch).dropna(how='all')
                    df_in.columns = df_in.columns.str.strip()
                    contratos_existentes = cargar_contratos()
                    errores = []
                    
                    # 1. Capa de Validación Atómica
                    for i, r in df_in.iterrows():
                        f_xl = i + 2 # Fila Excel aproximada considerando cabecera
                        
                        # Revisar obligatorios
                        c_obs = ['Empresa', 'Nombre', 'Canon', 'Moneda', 'Clase_Activo', 'Inicio', 'Fin', 'Tasa Anual %', 'Frecuencia_Pago']
                        for c in c_obs:
                            if c not in r or pd.isna(r[c]) or str(r[c]).strip() == "":
                                errores.append(f"Fila {f_xl}: Campo obligatorio '{c}' está vacío.")
                                
                        if str(r.get('Frecuencia_Pago', '')).strip() not in nombres_frec:
                            errores.append(f"Fila {f_xl}: Frecuencia_Pago '{r.get('Frecuencia_Pago')}' inválida. Opciones: {', '.join(nombres_frec)}")
                            
                        if any(f"Fila {f_xl}:" in e for e in errores): continue
                        
                        # Revisar Numéricos
                        try:
                            float(r['Canon'])
                            float(r['Tasa Anual %'])
                        except ValueError:
                            errores.append(f"Fila {f_xl}: 'Canon' o 'Tasa Anual %' no son un número válido.")
                            
                        # Revisar Fechas lógicas
                        try:
                            f_i = pd.to_datetime(r['Inicio'])
                            f_f = pd.to_datetime(r['Fin'])
                            if f_i.year < 2000:
                                errores.append(f"Fila {f_xl}: Año de Inicio no puede ser menor a 2000.")
                            if f_i > f_f:
                                errores.append(f"Fila {f_xl}: La fecha de Inicio es mayor a la fecha de Fin.")
                        except Exception:
                            errores.append(f"Fila {f_xl}: Formato de fecha inválido en 'Inicio' o 'Fin'.")
                            
                    if errores:
                        st.error("🚨 **Carga Denegada:** Se encontraron errores en la plantilla. No se ha ingresado NINGÚN contrato a la base de datos.")
                        for e in errores:
                            st.warning(e)
                    else:
                        # 2. Inserción Segura
                        import time
                        t_ini_carga = time.time()
                        for _, r in df_in.iterrows():
                            emp = str(r['Empresa']).strip()
                            f_i = pd.to_datetime(r['Inicio'])
                            f_f = pd.to_datetime(r['Fin'])
                            diff = relativedelta(f_f, f_i)
                            p = diff.years * 12 + diff.months
                            if diff.days >= 15: p += 1
                            t_an = float(r['Tasa Anual %'])
                            mon = str(r['Moneda'])
                            
                            nuevo_masivo = {
                                "Codigo_Interno": generar_codigo_correlativo(emp, contratos_existentes), 
                                "Empresa": emp, "Clase_Activo": str(r['Clase_Activo']), "ID": str(r.get('ID', '')), 
                                "Proveedor": str(r.get('Proveedor', '')), "Cod1": "", "Cod2": "", "Nombre": str(r['Nombre']), 
                                "Moneda": mon, "Canon": float(r['Canon']), "Tasa": t_an/100, 
                                "Tasa_Mensual": pow(1+t_an/100, 1/12)-1, 
                                "Valor_Moneda_Inicio": obtener_tc_cache(mon, f_i), "Plazo": p, 
                                "Inicio": f_i.strftime('%Y-%m-%d'), "Fin": f_f.strftime('%Y-%m-%d'), 
                                "Estado": "Activo", 
                                "Tipo_Pago": str(r.get('Tipo_Pago', 'Vencido')).strip().capitalize(), 
                                "Costos_Directos": float(r.get('Costos_Directos', 0.0) if pd.notna(r.get('Costos_Directos')) else 0.0), 
                                "Pagos_Anticipados": float(r.get('Pagos_Anticipados', 0.0) if pd.notna(r.get('Pagos_Anticipados')) else 0.0), 
                                "Costos_Desmantelamiento": float(r.get('Costos_Desmantelamiento', 0.0) if pd.notna(r.get('Costos_Desmantelamiento')) else 0.0), 
                                "Incentivos": float(r.get('Incentivos', 0.0) if pd.notna(r.get('Incentivos')) else 0.0),
                                "Frecuencia_Pago": str(r['Frecuencia_Pago']).strip()
                            }
                            # Asegurar carga de campos custom si vienen en el excel
                            for cx in obtener_parametros('CAMPO_EXTRA'):
                                nuevo_masivo[cx] = str(r[cx]) if cx in r and pd.notna(r[cx]) else ''
                                
                            insertar_contrato(nuevo_masivo)
                            contratos_existentes.append(nuevo_masivo) # Corrección de trick para número de correlativo
                        if 'motor_cache' in st.session_state: st.session_state.motor_cache.clear()
                        t_fin_carga = time.time()
                        st.success(f"✅ **Contratos validados y cargados exitosamente**. Se han ingresado {len(df_in)} nuevas líneas al sistema.")
                        st.info(f"⏱️ Tiempo total de procesamiento y carga: {t_fin_carga - t_ini_carga:.2f} segundos.")
                except Exception as e:
                    st.error(f"Error procesando archivo: {e}")

    with t3:
        st.subheader("Base de Datos Completa: Contratos")
        lista_c = cargar_contratos()
        from db import cargar_remediciones_todas_agrupadas
        rems_grupos = cargar_remediciones_todas_agrupadas()
        for c in lista_c:
            c['Fin_Original'] = (pd.to_datetime(c['Inicio']) + pd.DateOffset(months=int(c['Plazo']))).strftime('%Y-%m-%d')
            rems = rems_grupos.get(c['Codigo_Interno'], [])
            if rems:
                last_rem = rems[-1]
                c['Nuevo_Canon'] = last_rem['Canon']
                c['Nueva_Tasa_Anual_%'] = last_rem['Tasa'] * 100
                c['Nuevo_Plazo'] = last_rem['Plazo']
                c['Nuevo_Fin'] = last_rem['Fin']
            else:
                c['Nuevo_Canon'] = None
                c['Nueva_Tasa_Anual_%'] = None
                c['Nuevo_Plazo'] = None
                c['Nuevo_Fin'] = None
                
        df_contratos = pd.DataFrame(lista_c)
        if not df_contratos.empty:
            # Calcular Estado de Vigencia
            ahoy = date.today()
            df_contratos.insert(1, "Estado Vigencia", df_contratos['Fin'].apply(
                lambda x: '🚨 Vencido' if pd.to_datetime(x).date() < ahoy else '🟢 Vigente'
            ))
            
            # Filtrar solo columnas oficiales y campos extras activos (evitar zombies de SQLite)
            columnas_base = ['Codigo_Interno', 'Empresa', 'Clase_Activo', 'ID', 'Proveedor', 'Nombre', 'Moneda', 'Canon', 'Tasa', 'Tasa_Mensual', 'Valor_Moneda_Inicio', 'Plazo', 'Inicio', 'Fin', 'Fin_Original', 'Estado', 'Fecha_Baja', 'Ajuste_ROU', 'Tipo_Pago', 'Fecha_Remedicion', 'Frecuencia_Pago']
            permitidas = ['Estado Vigencia'] + columnas_base + obtener_parametros('CAMPO_EXTRA') + ['Nuevo_Canon', 'Nueva_Tasa_Anual_%', 'Nuevo_Plazo', 'Nuevo_Fin']
            cols_to_keep = [col for col in permitidas if col in df_contratos.columns]
            
            df_display = df_contratos[cols_to_keep]
            if 'Tasa' in df_display.columns:
                df_display = df_display.rename(columns={'Tasa': 'Tasa Anual %'})
            st.dataframe(df_display, use_container_width=True)
            st.download_button("Exportar Contratos (Excel)", to_excel(df_display), "Base_Contratos_Completa.xlsx")
        else:
            st.info("No hay contratos cargados.")

    with t4:
        st.subheader("Modificación de Contrato Existente")
        contratos_activos = [c for c in cargar_contratos() if c['Estado'] == 'Activo']
        if contratos_activos:
            mapa_c = {f"{c['Codigo_Interno']} - {c['Nombre']}": c for c in contratos_activos}
            sel = st.selectbox("Seleccione el Contrato a Modificar", list(mapa_c.keys()))
            c_sel = mapa_c[sel]
            
            st.write(f"**Condiciones Actuales:** Canon: {c_sel['Canon']} | Tasa Anual: {c_sel['Tasa']*100:.2f}% | Plazo Actual: {c_sel['Plazo']} meses")
            
            with st.form("f_rem"):
                st.write("Determine las nuevas condiciones de renovación o alteración del contrato.")
                c1, c2, c3 = st.columns(3)
                n_can = c1.number_input("Nuevo Canon", value=float(c_sel['Canon']), format="%.4f")
                n_tas = c2.number_input("Nueva Tasa Anual %", value=float(c_sel['Tasa']*100))
                n_fin = c3.date_input("Nueva Fecha Fin", value=pd.to_datetime(c_sel['Fin']), min_value=date(1900, 1, 1), max_value=date(2100, 12, 31))
                
                f_rem = st.date_input("Fecha Efectiva de Registro (Modificación)", value=date.today(), min_value=date(1900, 1, 1), max_value=date(2100, 12, 31))
                
                if st.form_submit_button("Aplicar Modificación"):
                    f_i = pd.to_datetime(c_sel['Inicio'])
                    f_rem_dt = pd.to_datetime(f_rem)
                    
                    if f_rem_dt <= f_i:
                        st.error("La fecha de modificación debe ser estrictamente posterior a la fecha de inicio original.")
                        st.stop()
                        
                    # Simulamos el contrato hasta la fecha de remedición para obtener los saldos de corte
                    tab_old, vp_old, rou_old = obtener_motor_financiero(c_sel)
                    
                    past_tab = tab_old[tab_old['Fecha'] < f_rem_dt]
                    if past_tab.empty:
                        st.error("La fecha seleccionada no permite capturar saldos históricos.")
                        st.stop()
                        
                    old_pasivo_orig = past_tab.iloc[-1]['S_Fin_Orig']
                    old_amort_acum_orig = past_tab['Dep_Orig'].sum()
                    old_rou_net_orig = rou_old - old_amort_acum_orig
                    
                    tc_ini = float(c_sel['Valor_Moneda_Inicio']) if float(c_sel['Valor_Moneda_Inicio']) > 0 else 1.0
                    tc_f_rem = obtener_tc_cache(c_sel['Moneda'], f_rem_dt)
                    if tc_f_rem == 0: tc_f_rem = 1.0
                    
                    diff = relativedelta(n_fin, f_rem_dt)
                    n_p = diff.years * 12 + diff.months
                    if diff.days >= 15: n_p += 1
                    t_m = pow(1+n_tas/100, 1/12)-1
                    
                    # Lógica de Reducción de Alcance (Terminación Parcial)
                    baja_pasivo_uf = 0.0
                    baja_rou_uf = 0.0
                    pl_efecto_clp = 0.0
                    
                    f_fin_old = pd.to_datetime(c_sel['Fin'])
                    if n_fin < f_fin_old.date():
                        diff_old = relativedelta(f_fin_old, f_rem_dt)
                        n_p_old = diff_old.years * 12 + diff_old.months
                        if diff_old.days >= 15: n_p_old += 1
                        
                        if n_p_old > 0:
                            p_dec = max(0.0, (n_p_old - n_p) / n_p_old)
                            baja_pasivo_uf = old_pasivo_orig * p_dec
                            baja_rou_uf = old_rou_net_orig * p_dec
                            # P&L in official currency (CLP) -> Pasivo decreases at current spot, ROU decreases at historical rate
                            pl_efecto_clp = (baja_pasivo_uf * tc_f_rem) - (baja_rou_uf * tc_ini)
                            
                            old_pasivo_orig -= baja_pasivo_uf
                            old_rou_net_orig -= baja_rou_uf
                    
                    # Cálculo del Puente (Ajuste ROU base funcional)
                    ajuste_rou_uf = old_rou_net_orig * (tc_ini / tc_f_rem) - old_pasivo_orig
                    
                    from db import insertar_remedicion, actualizar_contrato_remedicion
                    
                    # 1. Registrar el evento en el historial de remediciones con efectos P&L
                    insertar_remedicion(c_sel['Codigo_Interno'], f_rem.strftime('%Y-%m-%d'), n_can, n_tas/100, t_m, n_fin.strftime('%Y-%m-%d'), n_p, ajuste_rou_uf, baja_pasivo_uf, baja_rou_uf, pl_efecto_clp)
                    
                    # 2. Actualizar la cabecera del contrato original con la nueva fecha de madurez, para que los filtros de app sigan viéndolo activo hasta n_fin
                    # Importante: No machacamos Inicio, Canon base ni VP original. Solo los parámetros que avisan cuándo termina.
                    # CRITICO: El Plazo se mantiene ESTRICTAMENTE igual para no corromper la matemática del VP original.
                    actualizar_contrato_remedicion(c_sel['Codigo_Interno'], c_sel['Canon'], c_sel['Tasa'], c_sel['Tasa_Mensual'], n_fin.strftime('%Y-%m-%d'), c_sel['Plazo'], f_rem.strftime('%Y-%m-%d'))
                    
                    st.success(f"¡Contrato {c_sel['Codigo_Interno']} modificado exitosamente! (Se agregó el tramo de modificación a su flujo histórico)")
                    # Limpiar estado del motor financiero para forzar re-cálculo
                    if 'motor_cache' in st.session_state: st.session_state.motor_cache.clear()
                    st.cache_data.clear()
                    st.rerun()
        else:
            st.info("No hay contratos activos.")
            
    with t5:
        st.subheader("Baja Anticipada de Contrato")
        if 'contratos_activos' in locals() and contratos_activos:
            sel_b = st.selectbox("Seleccione Contrato para Baja", list(mapa_c.keys()), key="sbaja")
            c_baja = mapa_c[sel_b]
            f_baja = st.date_input("Fecha Efectiva de Baja", value=pd.to_datetime(c_baja['Fin']), min_value=date(1900, 1, 1), max_value=date(2100, 12, 31))
            if st.button("Procesar Baja Definitiva"):
                dar_baja_contrato(c_baja['Codigo_Interno'], f_baja.strftime('%Y-%m-%d'))
                if 'motor_cache' in st.session_state: st.session_state.motor_cache.clear()
                st.session_state.success_msg = f"Contrato dado de baja exitosamente en la fecha {f_baja}"
                st.rerun()

    with t6:
        st.subheader("Modificación Masiva de Contratos por Excel")
        st.write("Descargue la plantilla, complete únicamente los valores que desea modificar en las columnas 'Nuevo' y asigne una Fecha Efectiva, luego suba el archivo.")
        
        contratos_activos = [c for c in cargar_contratos() if c['Estado'] == 'Activo']
        if contratos_activos:
            df_plantilla = pd.DataFrame(contratos_activos)
            ahoy = date.today()
            df_plantilla['Estado Vigencia'] = df_plantilla['Fin'].apply(lambda x: '🚨 Vencido' if pd.to_datetime(x).date() < ahoy else '🟢 Vigente')
            
            df_plantilla['Tasa_Actual_%'] = df_plantilla['Tasa'] * 100
            df_plantilla.rename(columns={'Canon': 'Canon_Actual', 'Fin': 'Fecha_Fin_Actual'}, inplace=True)
            
            cols_solicitadas = ['Codigo_Interno', 'Estado Vigencia', 'Empresa', 'Clase_Activo', 'ID', 'Proveedor', 'Nombre', 'Moneda', 'Canon_Actual', 'Tasa_Actual_%', 'Fecha_Fin_Actual']
            df_plantilla = df_plantilla[cols_solicitadas]
            
            df_plantilla['Nuevo_Canon'] = None
            df_plantilla['Nueva_Tasa_Anual_%'] = None
            df_plantilla['Nueva_Fecha_Fin'] = None
            df_plantilla['Fecha_Efectiva_Modificacion'] = None
            
            st.download_button("📥 Descargar Plantilla de Modificaciones", to_excel(df_plantilla), "plantilla_modificacion_masiva.xlsx")
            
            arch_mod = st.file_uploader("Subir Archivo de Modificaciones", type=["xlsx"], key="arch_mod_masivo")
            if arch_mod is not None:
                if st.button("Procesar Modificaciones Masivas", type="primary"):
                    try:
                        df_in = pd.read_excel(arch_mod)
                        df_cambios = df_in.dropna(subset=['Fecha_Efectiva_Modificacion']).copy()
                        
                        if df_cambios.empty:
                            st.warning("No se detectó ninguna fila con 'Fecha_Efectiva_Modificacion'. No se realizaron cambios.")
                        else:
                            mapa_c = {f"{c['Codigo_Interno']}": c for c in contratos_activos}
                            exitos = 0
                            errores = []
                            
                            from db import insertar_remedicion, actualizar_contrato_remedicion
                            
                            for idx, r in df_cambios.iterrows():
                                cid = str(r['Codigo_Interno']).strip()
                                if cid not in mapa_c:
                                    errores.append(f"Fila {idx+2}: Contrato {cid} inactivo o no existe.")
                                    continue
                                    
                                c_sel = mapa_c[cid]
                                f_rem = pd.to_datetime(r['Fecha_Efectiva_Modificacion'])
                                f_i = pd.to_datetime(c_sel['Inicio'])
                                
                                if f_rem <= f_i:
                                    errores.append(f"Contrato {cid}: La fecha efectiva ({f_rem.date()}) no puede ser anterior al inicio ({f_i.date()}).")
                                    continue
                                    
                                n_can = float(r['Nuevo_Canon']) if pd.notna(r.get('Nuevo_Canon')) else float(c_sel['Canon'])
                                n_tas_pct = float(r['Nueva_Tasa_Anual_%']) if pd.notna(r.get('Nueva_Tasa_Anual_%')) else float(c_sel['Tasa']*100)
                                n_fin = pd.to_datetime(r['Nueva_Fecha_Fin']) if pd.notna(r.get('Nueva_Fecha_Fin')) else pd.to_datetime(c_sel['Fin'])
                                
                                # Simulacion
                                tab_old, vp_old, rou_old = obtener_motor_financiero(c_sel)
                                past_tab = tab_old[tab_old['Fecha'] < f_rem]
                                
                                if past_tab.empty:
                                    errores.append(f"Contrato {cid}: No existen saldos históricos previos a {f_rem.date()}.")
                                    continue
                                    
                                old_pasivo_orig = past_tab.iloc[-1]['S_Fin_Orig']
                                old_amort_acum_orig = past_tab['Dep_Orig'].sum()
                                old_rou_net_orig = rou_old - old_amort_acum_orig
                                
                                tc_ini = float(c_sel['Valor_Moneda_Inicio']) if float(c_sel['Valor_Moneda_Inicio']) > 0 else 1.0
                                tc_f_rem = obtener_tc_cache(c_sel['Moneda'], f_rem)
                                if tc_f_rem == 0: tc_f_rem = 1.0
                                
                                diff = relativedelta(n_fin, f_rem)
                                n_p = diff.years * 12 + diff.months
                                if diff.days >= 15: n_p += 1
                                t_m = pow(1+n_tas_pct/100, 1/12)-1
                                
                                # Lógica de Reducción de Alcance Módulo Masivo
                                baja_pasivo_uf = 0.0
                                baja_rou_uf = 0.0
                                pl_efecto_clp = 0.0
                                
                                f_fin_old = pd.to_datetime(c_sel['Fin'])
                                if n_fin < f_fin_old:
                                    diff_old = relativedelta(f_fin_old, f_rem)
                                    n_p_old = diff_old.years * 12 + diff_old.months
                                    if diff_old.days >= 15: n_p_old += 1
                                    
                                    if n_p_old > 0:
                                        p_dec = max(0.0, (n_p_old - n_p) / n_p_old)
                                        baja_pasivo_uf = old_pasivo_orig * p_dec
                                        baja_rou_uf = old_rou_net_orig * p_dec
                                        pl_efecto_clp = (baja_pasivo_uf * tc_f_rem) - (baja_rou_uf * tc_ini)
                                        
                                        old_pasivo_orig -= baja_pasivo_uf
                                        old_rou_net_orig -= baja_rou_uf
                                
                                ajuste_rou_uf = old_rou_net_orig * (tc_ini / tc_f_rem) - old_pasivo_orig
                                
                                # Insertar
                                insertar_remedicion(cid, f_rem.strftime('%Y-%m-%d'), n_can, n_tas_pct/100, t_m, n_fin.strftime('%Y-%m-%d'), n_p, ajuste_rou_uf, baja_pasivo_uf, baja_rou_uf, pl_efecto_clp)
                                # CRITICO: El Plazo se mantiene igual (c_sel['Plazo']) para no estropear la matematica original
                                actualizar_contrato_remedicion(cid, c_sel['Canon'], c_sel['Tasa'], c_sel['Tasa_Mensual'], n_fin.strftime('%Y-%m-%d'), c_sel['Plazo'], f_rem.strftime('%Y-%m-%d'))
                                exitos += 1
                                
                            if 'motor_cache' in st.session_state: st.session_state.motor_cache.clear()
                            st.cache_data.clear()
                            
                            if errores:
                                st.error("Se procesaron algunas modificaciones con advertencias.")
                                for e in errores: st.warning(e)
                            
                            if exitos > 0:
                                st.session_state.success_msg = f"¡Remedición Masiva Completada! Se modificaron {exitos} contratos exitosamente."
                                st.rerun()
                                
                    except Exception as e:
                        st.error(f"Error procesando el archivo: {e}")
        else:
            st.info("No hay contratos activos para descargar.")

def modulo_vencimientos():
    st.header("📝 Notas a los Estados Financieros")
    
    st.write("Esta nota revela el **Riesgo de Liquidez** al mostrar la salida futura bruta nominal de caja, clasificada por rangos de vencimiento. No incluye el descuento financiero, por lo tanto la suma difiere del Pasivo en Balance.")
    
    c1, c2, c3 = st.columns(3)
    m_nom = c1.selectbox("Mes de Cierre", MESES_LISTA, key="n_mes")
    a = c2.number_input("Año de Cierre", value=date.today().year, key="n_ano")
    
    df_c = pd.DataFrame(cargar_contratos())
    if df_c.empty:
        st.info("No hay contratos registrados.")
        return
        
    empresas = ["Todas"] + df_c['Empresa'].unique().tolist()
    emp_sel = c3.selectbox("Empresa", empresas, key="n_emp")
    
    c_btn1, c_btn2 = st.columns(2)
    with c_btn1:
        btn_no_desc = st.button("Generar Pasivos No Descontados", type="primary")
    with c_btn2:
        btn_desc = st.button("Generar Pasivos Descontados", type="primary")
        
    if btn_no_desc or btn_desc:
        es_desc = btn_desc
        f_t = pd.to_datetime(date(a, MESES_LISTA.index(m_nom)+1, 1)) + relativedelta(day=31)
        res = []
        
        # PRECARGAR REMEDICIONES PARA EVITAR QUERY N+1 (Pasa de 80s a <1s)
        from db import cargar_remediciones_todas_agrupadas
        rems_grupos = cargar_remediciones_todas_agrupadas()
        
        for _, c in df_c.iterrows():
            if emp_sel != "Todas" and c['Empresa'] != emp_sel: continue
            if f_t < pd.to_datetime(c['Inicio']).replace(day=1): continue
            
            # Si el contrato fue dado de baja antes o durante el mes de reporte, saltarlo
            if c.get('Fecha_Baja') and c['Estado'] == 'Baja':
                f_baja = pd.to_datetime(c['Fecha_Baja'])
                if f_baja <= f_t: continue
                
            tab, _, _ = obtener_motor_financiero(c, rems=rems_grupos.get(c['Codigo_Interno'], []))
            if tab.empty or 'Fecha' not in tab.columns: continue
            # Solo los flujos estrictamente futuros al cierre
            futuros = tab[tab['Fecha'] > f_t]
            if futuros.empty: continue
            
            # Para que cuadre perfectamente con el saldo contable en "Pasivos Descontados",
            # el monto a distribuir en los buckets debe ser la cuota de capital real deducida u o calculada 
            # desde la amortización del pasivo original.
            # Recordar que Capital_Mes = Pago_Orig - Int_Orig  (Lo que disminuye S_Fin_Orig)
            # Para la primera cuota corriente del mes siguiente, S_Ini_Orig es el Saldo del Pasivo Total de HOY.
            
            # saldo_remanente_hoy es la fotografia real de pasivo en el balance a f_t
            saldo_remanente_hoy = tab[tab['Fecha'] <= f_t].iloc[-1]['S_Fin_Orig'] if not tab[tab['Fecha'] <= f_t].empty else tab.iloc[0]['S_Ini_Orig']
            
            tc = obtener_tc_cache(c['Moneda'], f_t)
            
            ultima_fecha = futuros.iloc[-1]['Fecha']
            
            if es_desc:
                # La disminución del pasivo que genera esta cuota es: (S_Ini_Orig - S_Fin_Orig).
                futuros['Capital'] = futuros['S_Ini_Orig'] - futuros['S_Fin_Orig']
                # La ultima fila debe agregar todo el S_Fin_Orig remanente a su capital
                futuros.iloc[-1, futuros.columns.get_loc('Capital')] += futuros.iloc[-1]['S_Fin_Orig']
                suma_distribuida_orig = futuros['Capital'].sum()
                futuros['Monto_Bruto'] = futuros['Capital'] * tc
            else:
                futuros['Monto_Bruto'] = futuros['Pago_Orig'] * tc
                suma_distribuida_orig = 0
                
            limite_12m = f_t + relativedelta(months=12)
            
            # Vectorized assignment of es_corriente
            dias_pago = (futuros['Fecha'] - f_t).dt.days
            futuros['es_corriente'] = (dias_pago <= 90) | (futuros['Fecha'] <= limite_12_dash if 'limite_12_dash' in locals() else futuros['Fecha'] <= limite_12m)
            
            # Vectorized Assignment of Buckets
            bins = [-float('inf'), 90, 1095, 2555, float('inf')] # 90 days, 3 years, 7 years, >7 years
            # Corriente > 90 will be handled by boolean logic overrides later to mix date and day rules
            
            # Fallback direct assignment via map/apply for complex overlapping logic (small dataframe per contract usually)
            def assign_bucket(row):
                d = (row['Fecha'] - f_t).days
                if row['es_corriente']:
                    if d <= 90: return ('90 días', 1)
                    else: return ('90 días a 1 año', 2)
                else:
                    if d <= 1095: return ('2 a 3 años', 3)
                    elif d <= 2555: return ('4 a 7 años', 4)
                    else: return ('Más de 7 años', 5)
            
            buckets_ord = futuros.apply(assign_bucket, axis=1)
            futuros['Bucket'] = [b[0] for b in buckets_ord]
            futuros['Orden'] = [b[1] for b in buckets_ord]
            
            # Extraer dict para procesar el residuo contable final
            cuotas_temp = futuros[['Fecha', 'Bucket', 'Orden', 'Monto_Bruto', 'es_corriente']].to_dict('records')
            # Fix naming para compatibilidad
            for d in cuotas_temp: 
                d['Monto'] = d.pop('Monto_Bruto')
                d['ID_Contrato'] = c['Codigo_Interno']
                d['Nombre'] = c['Nombre']
                d['Clase_Activo'] = c['Clase_Activo']
            
            if es_desc and cuotas_temp:
                diferencia_residual = saldo_remanente_hoy - suma_distribuida_orig
                if abs(diferencia_residual) > 0.01:
                    cuota_lejana = max(cuotas_temp, key=lambda x: x['Orden'])
                    cuota_lejana['Monto'] += (diferencia_residual * tc)
                    
            res.extend(cuotas_temp)
        
        if not res:
            st.session_state.venc_data = pd.DataFrame() # DataFrame vacío
        else:
            st.session_state.venc_data = pd.DataFrame(res)
            
        st.session_state.venc_params = {'m': m_nom, 'a': a, 'f_t': f_t, 'es_desc': es_desc}

    if 'venc_data' in st.session_state:
        df_res = st.session_state.venc_data
        m_saved = st.session_state.venc_params['m']
        a_saved = st.session_state.venc_params['a']
        f_t_saved = st.session_state.venc_params['f_t']
        es_desc_saved = st.session_state.venc_params.get('es_desc', False)
        tipo_lbl = "descontados" if es_desc_saved else "no descontados"
        
        if df_res.empty:
            st.warning(f"No hay flujos futuros a rendir (Pasivos {tipo_lbl} = 0).")
            return
            
        todas_cols = ['90 días', '90 días a 1 año', 'Total Corriente', '2 a 3 años', '4 a 7 años', 'Más de 7 años', 'Total No Corriente']
        
        t1, t2 = st.tabs(["Por Clase", "Detalle por Contrato individual"])
        
        with t1:
            st.subheader(f"Pasivos {tipo_lbl}")
            piv = df_res.groupby(['Clase_Activo', 'Bucket', 'Orden'])['Monto'].sum().unstack(['Bucket', 'Orden']).fillna(0)
            piv.columns = [col[0] for col in piv.columns.to_flat_index()]
            
            piv['Total Corriente'] = piv.get('90 días', 0) + piv.get('90 días a 1 año', 0)
            piv['Total No Corriente'] = piv.get('2 a 3 años', 0) + piv.get('4 a 7 años', 0) + piv.get('Más de 7 años', 0)
            
            cols_finales = [c for c in todas_cols if c in piv.columns]
            piv = piv[cols_finales]
            
            piv = piv / 1000
            piv.loc['Total'] = piv.sum()
            
            st.write(f"**Detalle al {f_t_saved.strftime('%d-%m-%Y')} (En M$)**")
            st.dataframe(piv.style.format(precision=0, thousands="."))
            st.download_button("Exportar Formato (Excel)", to_excel(piv), f"Nota_Vencimientos_{m_saved}_{a_saved}.xlsx")
    
        with t2:
            st.subheader(f"Pasivos {tipo_lbl} (Detallado)")
            piv2 = df_res.groupby(['ID_Contrato', 'Nombre', 'Clase_Activo', 'Bucket', 'Orden'])['Monto'].sum().unstack(['Bucket', 'Orden']).fillna(0)
            piv2.columns = [col[0] for col in piv2.columns.to_flat_index()]
            
            piv2['Total Corriente'] = piv2.get('90 días', 0) + piv2.get('90 días a 1 año', 0)
            piv2['Total No Corriente'] = piv2.get('2 a 3 años', 0) + piv2.get('4 a 7 años', 0) + piv2.get('Más de 7 años', 0)
            
            cols_finales2 = [c for c in todas_cols if c in piv2.columns]
            piv2 = piv2[cols_finales2]
            
            piv2 = piv2 / 1000
            st.write(f"**Detalle Contractual al {f_t_saved.strftime('%d-%m-%Y')} (En M$)**")
            st.dataframe(piv2.style.format(precision=0, thousands="."))
            st.download_button("Exportar Detalle (Excel)", to_excel(piv2), f"Nota_Venc_Detalle_{m_saved}_{a_saved}.xlsx")


def modulo_auditoria():
    st.header("🔍 Auditoría y Transparencia")
    t1, t2 = st.tabs(["Fórmulas y Criterios", "Descarga de Datos Crudos"])
    
    with t1:
        st.subheader("Criterios Matemáticos de Cálculo IFRS 16")
        st.markdown('''
        **Motor Financiero V21.0 - Estándar IFRS 16**
        
        1. **Conversión de Tasa de Interés (Tasa Efectiva Mensual)** *(Ref. NIIF 16 Párraf. 26)*
        Se utiliza la fórmula de interés compuesto para hallar la tasa mensual equivalente a partir del input anual:
        `Tasa_Mensual = (1 + Tasa_Anual) ^ (1/12) - 1`
        *(Nota: Si se desea una validación lineal con calculadoras Excel estándar, se requiere proveer la tasa Nominal y no la Efectiva)*
        
        2. **Cálculo de Valor Presente (VP) - Pagos Vencidos** *(Ref. NIIF 16 Párraf. 26 - Medición inicial del pasivo)*
        `VP = Canon * [1 - (1 + Tasa_Mensual)^(-Plazo)] / Tasa_Mensual`
        
        3. **Cálculo de Valor Presente (VP) - Pagos Anticipados** *(Ref. NIIF 16 Párraf. 26 - Medición inicial del pasivo)*
        `VP = Canon * [1 - (1 + Tasa_Mensual)^(-Plazo)] / Tasa_Mensual * (1 + Tasa_Mensual)`
        
        4. **Cálculo del Activo por Derecho de Uso (ROU Inicial)** *(Ref. NIIF 16 Párraf. 24 - Medición inicial del activo)*
        `ROU = VP + Costos_Directos_Iniciales + Pagos_Anticipados_Extra + Costos_Desmantelamiento - Incentivos`
        
        5. **Amortización Mensual (Línea Recta)** *(Ref. NIIF 16 Párraf. 31 y 32 - Depreciación)*
        `Amortización = ROU_Inicial / Plazo`
        
        6. **Devengo de Intereses (Interés Efectivo)** *(Ref. NIIF 16 Párraf. 36 y 37 - Medición posterior del pasivo)*
        `Interés_Mes = Saldo_Inicial_Capital * Tasa_Mensual` (Para vencidos)
        `Interés_Mes = (Saldo_Inicial_Capital - Canon) * Tasa_Mensual` (Para anticipados)
        
        7. **Cálculo de Pasivos No Descontados (Análisis de Vencimientos)** *(Ref. NIIF 16 Párraf. 58 y NIIF 7 Párraf. 39(a))*
        Se utiliza para la nota obligatoria de riesgo de liquidez NIIF 7. Corresponde al sumatorio estricto de todos los desembolsos nominales brutos (`Canon * Tipo_Cambio_Cierre`) cuyas fechas de pago sean posteriores a la fecha de reporte, sin aplicar la tasa de descuento.
        `Pasivo_No_Descontado_Bucket_A = SUMA(Cánones_Futuros_Rango_A) * TC_Cierre`
        ''')
    
    with t2:
        st.subheader("Extracción de Data en Bruto")
        st.info("Descarga la base de datos subyacente de contratos activos para trazar los cálculos uno a uno en Excel.")
        df = pd.DataFrame(cargar_contratos())
        if not df.empty:
            st.download_button(
                label="📥 Descargar Base Completa Contratos",
                data=to_excel(df),
                file_name="Auditoria_Contratos_Bruto.xlsx"
            )

def _render_integracion_erp():
    st.subheader("Configuración de Integración ERP Local/Cloud")
    from db import guardar_credencial_erp, leer_credencial_erp, obtener_erp_activo
    
    erp_list = ["", "Odoo", "SAP ERP (OData/BAPI)", "Microsoft Dynamics 365", "Oracle NetSuite"]
    
    erp_activo_actual = obtener_erp_activo()
    st.info(f"ERP Activo Configurado Actualmente: **{erp_activo_actual if erp_activo_actual else 'Ninguno'}**")
    
    sel_erp = st.selectbox("Seleccione ERP a configurar", erp_list)
    
    if sel_erp == "Odoo":
        st.write("#### Credenciales de Odoo (XML-RPC)")
        curr_od = leer_credencial_erp("Odoo")
        with st.form("f_odoo"):
            o_url = st.text_input("URL del Servidor", value=curr_od['secretos'].get('url', 'https://miempresa.odoo.com'))
            o_db = st.text_input("Base de Datos", value=curr_od['secretos'].get('db', 'mi_bd'))
            o_user = st.text_input("Usuario (Email)", value=curr_od['secretos'].get('user', 'admin@empresa.com'))
            o_pass = st.text_input("Contraseña o API Key", type="password", value=curr_od['secretos'].get('pass', ''))
            o_act = st.checkbox("Activar esta Integración", value=curr_od['activo'])
            if st.form_submit_button("Guardar Configuración Odoo"):
                guardar_credencial_erp("Odoo", o_act, {'url': o_url, 'db': o_db, 'user': o_user, 'pass': o_pass})
                st.session_state.success_msg = "Credenciales de Odoo guardadas."
                st.rerun()
                
    elif sel_erp == "SAP ERP (OData/BAPI)":
        st.write("#### Credenciales de SAP S/4HANA / Business One")
        curr_sap = leer_credencial_erp("SAP ERP (OData/BAPI)")
        with st.form("f_sap"):
            s_url = st.text_input("URL Base (OData / Service Layer)", value=curr_sap['secretos'].get('url', 'https://sap.local:50000/'))
            s_cli = st.text_input("Client (Mandante)", value=curr_sap['secretos'].get('client', '100'))
            s_user = st.text_input("Usuario SAP", value=curr_sap['secretos'].get('user', 'API_USER'))
            s_pass = st.text_input("Contraseña", type="password", value=curr_sap['secretos'].get('pass', ''))
            s_act = st.checkbox("Activar esta Integración", value=curr_sap['activo'])
            if st.form_submit_button("Guardar Configuración SAP"):
                guardar_credencial_erp("SAP ERP (OData/BAPI)", s_act, {'url': s_url, 'client': s_cli, 'user': s_user, 'pass': s_pass})
                st.session_state.success_msg = "Credenciales de SAP guardadas."
                st.rerun()

    elif sel_erp == "Microsoft Dynamics 365":
        st.write("#### Credenciales Microsoft Dynamics (OAuth 2.0)")
        curr_ms = leer_credencial_erp("Microsoft Dynamics 365")
        with st.form("f_ms"):
            m_env = st.text_input("Entorno URL", value=curr_ms['secretos'].get('env', 'https://miempresa.operations.dynamics.com'))
            m_ten = st.text_input("Tenant ID", value=curr_ms['secretos'].get('tenant', 'xxx-xxx-xxx'))
            m_cli = st.text_input("Client ID", value=curr_ms['secretos'].get('client', ''))
            m_sec = st.text_input("Client Secret", type="password", value=curr_ms['secretos'].get('secret', ''))
            m_act = st.checkbox("Activar esta Integración", value=curr_ms['activo'])
            if st.form_submit_button("Guardar Configuración Microsoft"):
                guardar_credencial_erp("Microsoft Dynamics 365", m_act, {'env': m_env, 'tenant': m_ten, 'client': m_cli, 'secret': m_sec})
                st.session_state.success_msg = "Credenciales de Microsoft Dynamics guardadas."
                st.rerun()

    elif sel_erp == "Oracle NetSuite":
        st.write("#### Credenciales Oracle NetSuite (SuiteTalk REST/SOAP)")
        curr_ns = leer_credencial_erp("Oracle NetSuite")
        with st.form("f_ns"):
            n_acc = st.text_input("Account ID", value=curr_ns['secretos'].get('account', '1234567'))
            n_ckey = st.text_input("Consumer Key", value=curr_ns['secretos'].get('ckey', ''))
            n_csec = st.text_input("Consumer Secret", type="password", value=curr_ns['secretos'].get('csec', ''))
            n_tkey = st.text_input("Token ID", value=curr_ns['secretos'].get('tkey', ''))
            n_tsec = st.text_input("Token Secret", type="password", value=curr_ns['secretos'].get('tsec', ''))
            n_act = st.checkbox("Activar esta Integración", value=curr_ns['activo'])
            if st.form_submit_button("Guardar Configuración NetSuite"):
                guardar_credencial_erp("Oracle NetSuite", n_act, {'account': n_acc, 'ckey': n_ckey, 'csec': n_csec, 'tkey': n_tkey, 'tsec': n_tsec})
                st.session_state.success_msg = "Credenciales de NetSuite guardadas."
                st.rerun()

def modulo_configuracion():
    st.header("⚙️ Configuración del Sistema")
    
    if st.session_state.get('rol') == 'Ingeniero IT (Técnico)':
        t_it = st.tabs(["Integraciones ERP"])
        with t_it[0]:
            _render_integracion_erp()
        return
        
    opciones = ["Usuarios", "Empresas", "Monedas", "Campos Extra y Frecuencias", "Clases de Activo", "Cuentas Contables", "Integraciones ERP", "Bitácora de Auditoría", "Mantenimiento BD"]
    sel_tab = st.radio("Sección de Configuración", opciones, horizontal=True, key="config_tabs_radio")
    
    if sel_tab == "Usuarios":
        st.subheader("Gestión de Usuarios")
        c1, c2 = st.columns(2)
        n_user = c1.text_input("Nuevo Usuario")
        n_pass = c1.text_input("Contraseña", type="password")
        n_rol = c1.selectbox("Rol Estricto de Acceso", ["Administrador", "Analista Financiero (Editor)", "Auditor Ext. / Gerencia (Lector)", "Ingeniero IT (Técnico)"])
        if c1.button("Crear/Actualizar Usuario"):
            agregar_usuario(n_user, n_pass, n_rol)
            st.session_state.success_msg = f"Usuario {n_user} habilitado como {n_rol}."
            st.rerun()
            
        st.write("Usuarios Actuales en Sistema:")
        usr_dicts = obtener_usuarios() 
        usr_df = pd.DataFrame(usr_dicts)
        if not usr_df.empty:
            usr_df.rename(columns={'usuario': 'Usuario', 'rol': 'Perfil Acceso'}, inplace=True)
        st.dataframe(usr_df)
        
        st.subheader("Eliminar Usuario")
        lista_u = usr_df["Usuario"].tolist() if not usr_df.empty else []
        del_usr = st.selectbox("Seleccione Usuario a Eliminar", [""] + lista_u)
        if st.button("Eliminar", key="del_usr_btn") and del_usr != "":
            conn = conectar()
            conn.execute("DELETE FROM usuarios WHERE usuario=?", (del_usr,))
            conn.commit(); conn.close()
            st.session_state.success_msg = f"Usuario {del_usr} eliminado con éxito."
            st.rerun()
        
    elif sel_tab == "Empresas":
        st.subheader("Sociedades / Empresas")
        c1, c2 = st.columns(2)
        nueva_empresa = c1.text_input("Nombre de la Nueva Empresa")
        if c1.button("Agregar Empresa"):
            agregar_parametro('EMPRESA', nueva_empresa.strip())
            st.session_state.success_msg = "Empresa ingresada con éxito."
            st.rerun()
            
        st.write("Empresas Activas:")
        emp_df = pd.DataFrame(obtener_parametros('EMPRESA'), columns=["Empresa_Registrada"])
        st.dataframe(emp_df)
        
        st.subheader("Modificar / Eliminar Empresa")
        c3, c4 = st.columns(2)
        del_emp = c3.selectbox("Seleccione Empresa a Eliminar", [""] + emp_df["Empresa_Registrada"].tolist())
        if c3.button("Eliminar", key="del_emp_btn") and del_emp != "":
            eliminar_parametro('EMPRESA', del_emp)
            st.session_state.success_msg = f"Empresa '{del_emp}' eliminada."
            st.rerun()
            
        mod_emp_old = c4.selectbox("Seleccione Empresa a Renombrar", [""] + emp_df["Empresa_Registrada"].tolist())
        mod_emp_new = c4.text_input("Nuevo Nombre", key="new_n_emp")
        if c4.button("Renombrar Empresa") and mod_emp_old != "" and mod_emp_new.strip() != "":
            conn = conectar()
            # Renombrar en la config
            conn.execute("UPDATE config_params SET valor=? WHERE tipo='EMPRESA' AND valor=?", (mod_emp_new.strip(), mod_emp_old))
            # Actualizar todos los contratos asociados a esta empresa
            conn.execute("UPDATE contratos SET Empresa=? WHERE Empresa=?", (mod_emp_new.strip(), mod_emp_old))
            conn.commit(); conn.close()
            st.session_state.success_msg = "Empresa actualizada con éxito."
            st.rerun()
        
    elif sel_tab == "Monedas":
        st.subheader("Monedas Habilitadas en el Sistema")
        st.info("Agregue monedas para procesamientos futuros. No se podrán borrar si ya existen en algún contrato.")
        c1, c2 = st.columns(2)
        nueva_moneda = c1.text_input("Nueva Moneda (ej. JPY, GBP)")
        if c1.button("Agregar Moneda"):
            if nueva_moneda.strip() != "":
                agregar_parametro('MONEDA', nueva_moneda.strip().upper())
                st.session_state.success_msg = "Moneda ingresada con éxito."
                st.rerun()
            
        st.write("Monedas Activas:")
        mon_df = pd.DataFrame(obtener_parametros('MONEDA'), columns=["Moneda"])
        st.dataframe(mon_df)
        
        st.subheader("Modificar / Eliminar Moneda")
        c3, c4 = st.columns(2)
        del_mon = c3.selectbox("Seleccione Moneda a Eliminar", [""] + mon_df["Moneda"].tolist())
        if c3.button("Eliminar Moneda", key="del_mon_btn") and del_mon != "":
            rr = eliminar_parametro('MONEDA', del_mon)
            if rr:
                st.session_state.success_msg = f"Moneda '{del_mon}' eliminada."
                st.rerun()
            else:
                st.error(f"No se puede eliminar '{del_mon}'. Ya está en uso por contratos vigentes o históricos.")
                
        mod_mon_old = c4.selectbox("Seleccione Moneda a Renombrar", [""] + mon_df["Moneda"].tolist())
        mod_mon_new = c4.text_input("Nuevo Nombre M", key="new_n_mon")
        if c4.button("Renombrar Moneda") and mod_mon_old != "" and mod_mon_new.strip() != "":
            conn = conectar()
            conn.execute("UPDATE config_params SET valor=? WHERE tipo='MONEDA' AND valor=?", (mod_mon_new.strip().upper(), mod_mon_old))
            conn.execute("UPDATE contratos SET Moneda=? WHERE Moneda=?", (mod_mon_new.strip().upper(), mod_mon_old))
            conn.commit(); conn.close()
            st.session_state.success_msg = "Moneda actualizada con éxito."
            st.rerun()

    elif sel_tab == "Campos Extra y Frecuencias":
        st.subheader("Frecuencias de Pago Habilitadas")
        st.info("Formato requerido: NombreFrecuencia-MultiplicadorMeses. Ej: Cuatrimestral-4")
        c1_f, c2_f = st.columns(2)
        n_frec = c1_f.text_input("Nueva Frecuencia (Ej: Cuatrimestral-4)")
        if c1_f.button("Añadir Frecuencia"):
            if '-' in n_frec and n_frec.split('-')[1].isdigit():
                agregar_parametro('FRECUENCIA_PAGO', n_frec.strip())
                st.session_state.success_msg = "Frecuencia agregada exitosamente."
                st.rerun()
            else:
                st.error("Por favor respete el formato 'Nombre-Meses'.")
                
        f_df = pd.DataFrame(obtener_parametros('FRECUENCIA_PAGO'), columns=["Frecuencias (Nombre-Meses)"])
        st.dataframe(f_df)
        
        st.subheader("Eliminar Frecuencia de Pago")
        del_f = c2_f.selectbox("Frecuencia a Eliminar", [""] + f_df["Frecuencias (Nombre-Meses)"].tolist())
        if c2_f.button("Eliminar Frecuencia", key="del_frec_btn") and del_f != "":
            eliminar_parametro('FRECUENCIA_PAGO', del_f)
            st.session_state.success_msg = f"Frecuencia '{del_f}' eliminada."
            st.rerun()
                
        st.markdown("---")
        st.subheader("Campos Adicionales (Contratos)")
        st.info("Agregue columnas extra (ej. 'Patente', 'Centro Costo') al ingreso de Contratos. Se crearán dinámicamente en la Base de Datos.")
        c1, c2 = st.columns(2)
        nuevo_campo = c1.text_input("Nombre del Nuevo Campo")
        if c1.button("Crear Campo de Información"):
            if nuevo_campo.strip() != "":
                from db import invocar_columna_extra
                invocar_columna_extra(nuevo_campo.strip())
                st.session_state.success_msg = "Campo adicionado exitosamente."
                st.rerun()
        
        st.write("Campos Personalizados Creados:")
        cx_df = pd.DataFrame(obtener_parametros('CAMPO_EXTRA'), columns=["Campo Extra"])
        st.dataframe(cx_df)
        
        st.subheader("Eliminar Campo Extra")
        del_cx = st.selectbox("Seleccione Campo a Eliminar", [""] + cx_df["Campo Extra"].tolist())
        if st.button("Eliminar Campo", key="del_cx_btn") and del_cx != "":
            rr = eliminar_parametro('CAMPO_EXTRA', del_cx)
            if rr:
                st.session_state.success_msg = f"Campo '{del_cx}' eliminado/ocultado de la configuración."
                st.rerun()
            else:
                st.error(f"No se puede borrar '{del_cx}'. Existen contratos que ya tienen información diligenciada en esta columna.")

    elif sel_tab == "Clases de Activo":
        st.subheader("Clases de Activo")
        c1, c2 = st.columns(2)
        with c1.form("form_add_cls", clear_on_submit=True):
            nueva_clase = st.text_input("Ingresar Nueva Clase")
            if st.form_submit_button("Agregar Clase") and nueva_clase.strip() != "":
                agregar_parametro('CLASE_ACTIVO', nueva_clase.strip())
                st.session_state.success_msg = "Clase de activo ingresada con éxito."
                st.rerun()
            
        st.write("Clases Activas:")
        cls_df = pd.DataFrame(obtener_parametros('CLASE_ACTIVO'), columns=["Clase_Registrada"])
        st.dataframe(cls_df)
        
        st.subheader("Modificar / Eliminar Clase")
        c3, c4 = st.columns(2)
        with c3.form("form_del_cls"):
            del_cls = st.selectbox("Seleccione Clase a Eliminar", [""] + cls_df["Clase_Registrada"].tolist())
            if st.form_submit_button("Eliminar") and del_cls != "":
                eliminar_parametro('CLASE_ACTIVO', del_cls)
                st.session_state.success_msg = f"Clase '{del_cls}' eliminada."
                st.rerun()
            
        with c4.form("form_ren_cls", clear_on_submit=True):
            mod_cls_old = st.selectbox("Seleccione Clase a Renombrar", [""] + cls_df["Clase_Registrada"].tolist())
            mod_cls_new = st.text_input("Nuevo Nombre", key="new_n_cls")
            if st.form_submit_button("Renombrar Clase") and mod_cls_old != "" and mod_cls_new.strip() != "":
                conn = conectar()
                conn.execute("UPDATE config_params SET valor=? WHERE tipo='CLASE_ACTIVO' AND valor=?", (mod_cls_new.strip(), mod_cls_old))
                conn.execute("UPDATE contratos SET Clase_Activo=? WHERE Clase_Activo=?", (mod_cls_new.strip(), mod_cls_old))
                conn.commit(); conn.close()
                st.session_state.success_msg = "Clase actualizada con éxito."
                st.rerun()

    elif sel_tab == "Cuentas Contables":
        st.subheader("Traductor Plan de Cuentas Automático")
        clist = [
            ('CUENTA_ROU', 'Activo por Derecho de Uso (ROU)'), 
            ('CUENTA_PASIVO', 'Pasivo por Arrendamiento'), 
            ('CUENTA_BCO_AJUSTE', 'Banco / Provisiones (Ajustes)'), 
            ('CUENTA_GASTO_AMORT', 'Gasto Amortización ROU'), 
            ('CUENTA_AMORT_ACUM', 'Amortización Acumulada ROU'), 
            ('CUENTA_GASTO_INT', 'Gasto Interés (Costo Fin.)'), 
            ('CUENTA_BANCO_PAGO', 'Banco Efectivo (Pago)'), 
            ('CUENTA_PERDIDA_TC', 'Pérdida por Tipo de Cambio'), 
            ('CUENTA_GANANCIA_TC', 'Ganancia por Tipo de Cambio')
        ]
        with st.form("fc_cuentas"):
            n_vals = []
            st.write("Ingrese el N° de Cuenta y Nombre de Cuenta para cada tipo.")
            for k, label_default in clist:
                col1, col2 = st.columns([1, 4])
                v_num_act = obtener_parametros(k + '_NUM')[0] if obtener_parametros(k + '_NUM') else "0000"
                v_nom_act = obtener_parametros(k + '_NOM')[0] if obtener_parametros(k + '_NOM') else label_default
                n_num = col1.text_input(f"N°", value=v_num_act, key=k+'_NUM')
                n_nom = col2.text_input(label_default, value=v_nom_act, key=k+'_NOM')
                n_vals.append((k + '_NUM', n_num))
                n_vals.append((k + '_NOM', n_nom))
            
            if st.form_submit_button("Actualizar y Guardar Plan de Cuentas"):
                conn = conectar()
                for k, v in n_vals:
                    conn.execute("DELETE FROM config_params WHERE tipo=?", (k,))
                    conn.execute("INSERT INTO config_params VALUES (?,?)", (k, v))
                conn.commit(); conn.close()
                st.session_state.success_msg = "Toda la contabilización del motor IFRS ha sido enrutada a este nuevo Plan de Cuentas."
                st.rerun()

    elif sel_tab == "Integraciones ERP":
        _render_integracion_erp()

    elif sel_tab == "Bitácora de Auditoría":
        st.subheader("Bitácora de Auditoría (Logs del Sistema)")
        st.info("Registre y trace todas las acciones operativas realizadas por los usuarios (creación/eliminación de contratos, modificaciones, etc).")
        
        c_tog, c_btn = st.columns([3, 1])
        estado_actual = is_audit_enabled()
        
        with c_tog:
            nuevo_estado = st.toggle("Activar Registro Automático de Acciones (Auditoría)", value=estado_actual)
            
        if nuevo_estado != estado_actual:
            if nuevo_estado:
                agregar_parametro('AUDIT_LOG_ENABLED', '1')
                st.session_state.success_msg = "Bitácora de Auditoría ACTIVADA. Todas las alteraciones en Contratos y Parámetros quedarán registradas."
            else:
                eliminar_parametro('AUDIT_LOG_ENABLED', '1')
                st.session_state.success_msg = "Bitácora DESACTIVADA."
            st.rerun()
            
        st.markdown("---")
        if c_btn.button("🔄 Refrescar Visor de Logs", type="primary", use_container_width=True):
            st.rerun()
            
        logs_df = obtener_logs()
        if logs_df.empty:
            st.warning("No hay registros en la Bitácora actualmente.")
        else:
            logs_df.rename(columns={'id':'ID_Log', 'fecha_hora':'Fecha/Hora', 'usuario':'Usuario', 'accion': 'Acción Ejecutada', 'entidad_id': 'Contrato / Entidad', 'detalles':'Detalle Extra'}, inplace=True)
            st.dataframe(logs_df, use_container_width=True)

    elif sel_tab == "Mantenimiento BD":
        st.subheader("Mantenimiento y Reseteo de Datos")
        st.warning("⚠️ **ATENCIÓN**: Las acciones de esta sección borrarán datos operativos del sistema. Asegúrese de haber exportado lo que necesite antes de proceder. La configuración (cuentas, usuarios, empresas) NO será afectada.")
        
        st.markdown("---")
        st.write("### 💱 1. Limpieza de Monedas (Tipos de Cambio)")
        st.write("Borra todo el historial de tipos de cambio de la base de datos para cargar un nuevo archivo desde cero.")
        if st.button("🗑️ Borrar Todos los Tipos de Cambio", type="primary", key="btn_limpiar_monedas"):
            limpiar_monedas()
            st.cache_data.clear()
            if 'motor_cache' in st.session_state: st.session_state.motor_cache.clear()
            st.success("✅ Historial de monedas eliminado por completo. Puede cargar un nuevo archivo en el módulo de Monedas.")
            
        st.markdown("---")
        st.write("### 📝 2. Limpieza de Contratos y Remediciones")
        st.write("Borra toda la cartera de contratos registrados y toda su historia de remediciones (trazabilidad).")
        if st.button("🚨 Borrar Todos los Contratos Registrados", type="primary", key="btn_limpiar_contratos"):
            limpiar_contratos()
            st.cache_data.clear()
            if 'motor_cache' in st.session_state: st.session_state.motor_cache.clear()
            st.success("✅ Base de datos de contratos vaciada. El sistema está listo para una carga masiva desde el módulo de Contratos.")

def resolver_tasa_implicita(vr, ca, canon, plazo_meses, vrng, oc):
    total_inversion = vr + ca
    if total_inversion <= 0 or canon <= 0 or plazo_meses <= 0:
        return 0.0
        
    vp_zero_rate = (canon * plazo_meses) + vrng + oc
    if vp_zero_rate < total_inversion:
        return -1.0
    
    # Búsqueda binaria para la tasa de interés TIR equivalente mensual
    low = 0.0
    high = 100.0 # Búsqueda hasta 10000% mensual para asegurar rango
    
    for _ in range(100):
        mid = (low + high) / 2.0
        if mid == 0:
            vp_canon = canon * plazo_meses
            vp_residual = vrng + oc
        else:
            vp_canon = canon * (1 - (1 + mid)**(-plazo_meses)) / mid
            vp_residual = (vrng + oc) / ((1 + mid)**plazo_meses)
            
        vp_total = vp_canon + vp_residual
        
        if vp_total > total_inversion:
            # Tasa muy chica, el VP está muy grande
            low = mid
        else:
            # Tasa muy grande, VP chico
            high = mid
            
    tasa_mensual = (low + high) / 2.0
    tasa_anual = ((1 + tasa_mensual)**12) - 1
    return tasa_anual * 100

def modulo_asistente_ibr():
    st.header("🧮 Asistente de calculos (tasas de contratos-Activo y pasivo ROU)")
    st.write("**Herramienta aislada de consulta - No afecta la base de datos operativa**")
    
    with st.expander("📖 Diagnóstico Inicial: Jerarquía IFRS 16 y Definiciones Básicas", expanded=False):
        st.markdown('''
        <div style="font-size: 0.9em; font-style: italic; color: #666; margin-bottom: 20px; line-height: 1.4;">
        <b>NIIF 16 - Párrafo 26: Medición Inicial del Pasivo por Arrendamiento</b><br>
        "En la fecha de comienzo, un arrendatario medirá el pasivo por arrendamiento al valor presente de los pagos por arrendamiento que no se hayan pagado en esa fecha. Los pagos por arrendamiento se descontarán usando la tasa de interés implícita en el arrendamiento, si esa tasa pudiera determinarse fácilmente. Si esa tasa no puede determinarse fácilmente, el arrendatario utilizará la tasa incremental por préstamo del arrendatario (IBR)."
        <br><br>
        <b>Definiciones:</b><br>
        • <b>Tasa Anual de interés Implícita</b>: Es la tasa de interés que iguala el valor presente de los pagos por arrendamiento y el valor residual no garantizado con el valor razonable del activo subyacente más cualquier costo directo inicial del arrendador.<br>
        • <b>Tasa Anual IBR (Incremental por préstamo)</b>: Es la tasa de interés que un arrendatario tendría que pagar por tomar prestado en un entorno económico similar.<br>
        • <b>Tasa Mensual Equivalente</b>: Es la tasa de descuento periódica derivada de la Tasa Anual mediante la fórmula de interés compuesto (raíz doceava). Esta tasa se aplica directamente sobre el saldo insoluto del pasivo en cada cierre mensual para registrar el gasto financiero del periodo.
        </div>
        ''', unsafe_allow_html=True)
        
    st.markdown("#### ¿Qué desea calcular?")
    tab_imp, tab_ibr, tab_men, tab_vp = st.tabs(["a) Determinar Tasa Anual Implícita", "b) Determinar Tasa Anual IBR", "c) Tasa de Interés Mensual", "d) Calculadora Valor Presente"])
            
    with tab_imp:
        st.subheader("Calculadora Tasa Implícita en el arrendamiento")
        st.info("💡 **Definición de Tasa Implícita:** Es la tasa de interés que iguala el valor presente de los pagos por arrendamiento y el valor residual no garantizado con el valor razonable del activo subyacente más cualquier costo directo inicial del arrendador.")
        col_m_imp, _ = st.columns(2)
        monedas_activas = obtener_parametros('MONEDA')
        if not monedas_activas: monedas_activas = ["UF", "CLP", "USD", "EUR"]
        moneda_global = col_m_imp.selectbox("Seleccione la Moneda del contrato", monedas_activas, key="moneda_imp")
        st.warning("⚠️ PARAMETROS BÁSICOS REQUERIDOS: Valor Razonable del bien, Valor residual al termino del contrato, Canon mensual y Plazo expresado en Meses. Si usted NO TIENE el monto económico preciso de alguno de estos parámetros comerciales, se recomienda utilizar de inmediato la TASA IBR para evitar errores matemáticos de cálculo financiero.")
        
        st.write("**A. Datos del Arrendador / Activo:**")
        c1, c2, c3 = st.columns(3)
        vr = c1.number_input(f"Valor Razonable (Mercado) del activo ({moneda_global})", min_value=0.0, step=100.0)
        vrng = c2.number_input(f"Valor Residual No Garantizado estim. ({moneda_global})", min_value=0.0, step=100.0)
        ca = c3.number_input(f"Costos Inic. Directos Arrendador ({moneda_global})", min_value=0.0, step=100.0)
        
        st.write("**B. Datos del Contrato de Arrendamiento (Flujos):**")
        c4, c5, c6 = st.columns(3)
        canon = c4.number_input(f"Canon Mensual / Cuota Regular ({moneda_global})", min_value=0.0, step=10.0)
        plazo_meses = c5.number_input("Plazo del Contrato Cierto (Meses)", min_value=0, max_value=600, value=0, step=1)
        oc = c6.number_input(f"Opción de Compra u Otros Flujos ({moneda_global})", min_value=0.0, step=100.0)
        
        if vr == 0 or canon == 0 or plazo_meses == 0:
            st.warning("**[!] ADVERTENCIA TÉCNICA:**\\nFaltan datos obligatorios para calcular matemáticamente la Tasa Implícita (Valor Razonable, Canon Mensual, Plazo). Según la NIIF 16 (párrafo 26), si no puede determinarse, el arrendatario utilizará su Tasa IBR.")
        else:
            if st.button("Calcular Tasa Implícita Matemáticamente", type="primary"):
                t_imp = resolver_tasa_implicita(vr, ca, canon, plazo_meses, vrng, oc)
                
                if t_imp < 0:
                    st.error("❌ **Error Matemático de Negocio:** La suma nominal de todos los flujos que retornarán al arrendador (Cuotas + Opción + Valor Residual) es inferior al Valor Razonable del activo que entregó. \nEsto indica que, con los datos ingresados, la rentabilidad del arrendador es negativa, lo cual es comercialmente irreal y no permite despejar una tasa matemática positiva.\\n\\n*Verifica que los ceros del Valor Razonable no estén sobrando o corrige el Canon/Valor Residual.*")
                else:
                    st.success(f"**Tasa Implícita Efectiva Anual Obtenida: {t_imp:.4f}%**")
                    st.latex(r"Valor\ Razonable + Costos\ Directos = \sum_{t=1}^{n} \frac{Canon}{(1+i)^t} + \frac{Valor\ Residual\ No\ Garantizado + Opci\acute{o}n\ de\ Compra}{(1+i)^n}")
                    
                    import docx
                    from docx.shared import Pt, Inches
                    import io
                    import urllib.request
                    import urllib.parse
                    
                    doc = docx.Document()
                    doc.add_heading("MEMORIA DE CÁLCULO - TASA IMPLÍCITA (IFRS 16)", 0)
                    doc.add_paragraph(f"Fecha: {datetime.datetime.now().strftime('%Y-%m-%d %H:%M:%S')}")
                    
                    doc.add_heading("1. Variables Ingresadas", level=1)
                    doc.add_paragraph(f"Moneda del Contrato: {moneda_global}")
                    doc.add_paragraph(f"Valor Razonable: {vr:,.2f}")
                    doc.add_paragraph(f"Costos Iniciales Arrendador: {ca:,.2f}")
                    doc.add_paragraph(f"Canon Mensual: {canon:,.2f}")
                    doc.add_paragraph(f"Plazo Contrato: {plazo_meses} meses")
                    doc.add_paragraph(f"Valor Residual No Garantizado: {vrng:,.2f}")
                    doc.add_paragraph(f"Opción de Compra u Otros Pagos: {oc:,.2f}")
                    
                    doc.add_heading("2. Fórmula de Cálculo", level=1)
                    doc.add_paragraph("La Tasa Implícita Efectiva es aquella que iguala el Valor Inversión del Arrendador con el Valor Presente (VP) de los flujos del contrato.")
                    try:
                        formula_str = r"\bg_white\dpi{150}\large Valor\ Razonable\ +\ Costos\ Directos\ =\ \sum_{t=1}^{n}\ \frac{Canon}{(1+i)^t}\ +\ \frac{Valor\ Residual\ No\ Garantizado\ +\ Opci\acute{o}n\ de\ Compra}{(1+i)^n}"
                        url = "https://latex.codecogs.com/png.latex?" + urllib.parse.quote(formula_str)
                        req = urllib.request.Request(url, headers={'User-Agent': 'Mozilla/5.0'})
                        with urllib.request.urlopen(req) as response:
                            img_data = response.read()
                        doc.add_picture(io.BytesIO(img_data), width=Inches(6.0))
                    except Exception:
                        doc.add_paragraph("Valor Inversión = Valor Razonable + Costos Iniciales")
                        doc.add_paragraph("VP Flujos = VP(Canon mensual durante el plazo) + VP(Valor Residual No Garantizado) + VP(Opción de Compra)")
                    doc.add_paragraph("Se utiliza un método iterativo (búsqueda binaria) para resolver matemáticamente la tasa TIR que cumple exactamente dicha igualdad.")
                    
                    doc.add_heading("3. Resultado Final", level=1)
                    p = doc.add_paragraph()
                    r = p.add_run(f"TASA IMPLÍCITA EFECTIVA ANUAL: {t_imp:.4f}%")
                    r.bold = True
                    r.font.size = Pt(14)
                    
                    output = io.BytesIO()
                    doc.save(output)
                    word_data = output.getvalue()
                    
                    st.download_button(
                        label="📄 Descargar Memoria de Cálculo Implícito (Word)",
                        data=word_data,
                        file_name=f"Tasa_Implicita_{datetime.datetime.now().strftime('%Y%m%d')}.docx",
                        mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
                    )

    with tab_ibr:
        st.subheader("Calculadora IBR Personalizado")
        st.info("💡 **Definición de Tasa IBR:** Es la tasa de interés anual que su empresa tendría que pagar al banco o acreedor por tomar prestado, en un plazo similar y con una seguridad similar, los fondos necesarios para obtener un activo de valor similar al derecho de uso en un entorno económico similar.")
        col_m_ibr, _ = st.columns(2)
        monedas_activas = obtener_parametros('MONEDA')
        if not monedas_activas: monedas_activas = ["UF", "CLP", "USD", "EUR"]
        moneda_global = col_m_ibr.selectbox("Seleccione la Moneda del contrato", monedas_activas, key="moneda_ibr")
        
        with st.expander("📖 Enlaces de Referencia de Tasas (BCU, Treasury, etc.)", expanded=True):
            st.markdown('''
            - **Tasa IBR**: Incremental Borrowing Rate (Tasa incremental por préstamos).
            - **BCU (UF)** / **BCP (Pesos)**: [bcentral.cl](https://bcentral.cl)
            - **Treasury Yield (USD)**: [treasurydirect.gov](https://treasurydirect.gov)
            - **Bunds (EUR - Alemania)**: [deutsche-finanzagentur.de](https://deutsche-finanzagentur.de)
            - **Tasa Implícita**: Aquella tasa de descuento que iguala el VP de los Flujos de Efectivo con el Valor Razonable del Activo + Costos Directos.
            ''')
            
        st.markdown("**A. Determinación de la Tasa Base Libre de Riesgo**")
        st.caption("**El Nuevo Contrato**: Aquí se evalúan las condiciones macroeconómicas puras del contrato (Moneda y Plazo), para obtener la tasa inicial libre de riesgo aplicable en el mercado.")
        
        c_cuotas, c_tasa = st.columns(2)
        plazo_meses_ibr = c_cuotas.number_input("Número de Cuotas (Plazo en Meses)", min_value=1, max_value=600, value=60, step=1)
        
        plazo_anios = plazo_meses_ibr / 12.0
        if moneda_global == "UF":
            desc_tasa = f"Ingrese Tasa BCU a {plazo_anios:.1f} años (% Anual)"
        elif moneda_global == "CLP":
            desc_tasa = f"Ingrese Tasa BCP a {plazo_anios:.1f} años (% Anual)"
        elif moneda_global == "USD":
            desc_tasa = f"Treasury Yield a {plazo_anios:.1f} años (% Anual)"
        else:
            desc_tasa = f"Bunds (Alemania) a {plazo_anios:.1f} años (% Anual)"
            
        tasa_base = c_tasa.number_input(desc_tasa, min_value=0.0, value=3.0, step=0.1)
        
        st.markdown("**B. Determinación del Spread de Riesgo Ponderado**")
        st.caption("**El Perfil de la Empresa**: Aquí se evalúa el historial crediticio corporativo para medir el 'sobrecargo' por riesgo que los acreedores exigen comercialmente a la compañía.")
        if moneda_global == "UF":
            nombre_instrumento = "Bono del Banco Central en UF - BCU"
        elif moneda_global == "CLP":
            nombre_instrumento = "Bono del Banco Central en Pesos - BCP"
        elif moneda_global == "USD":
            nombre_instrumento = "US Treasury Yield"
        else:
            nombre_instrumento = "Bunds - Bonos Soberanos de Alemania"
            
        st.info(f'''💡 **Guía de Carga de Deudas Históricas:**
- **Tasa Deuda Asignada (% Anual):** La tasa de interés cobrada por el banco en ese crédito comercial específico.
- **Tasa Ref. Histórica (% Anual):** La Tasa Libre de Riesgo (**extrayendo el {nombre_instrumento}**) que existía *exactamente en la misma fecha* que se firmó el crédito, **y por el mismo plazo original de esa deuda**.

*El sistema restará ambas automáticamente para aislar el sobrecargo originario por Riesgo de Empresa.*''')
        
        num_deudas = st.number_input("¿Cuántas deudas vigentes posee para determinar el Spread de Riesgo?", min_value=0, max_value=10, value=1, step=1)
        
        deudas = []
        suma_capital = 0.0
        suma_ponderada = 0.0
        
        if num_deudas > 0:
            st.write("Ingrese el detalle de cada deuda:")
            for i in range(num_deudas):
                cols = st.columns(3)
                cap = cols[0].number_input(f"Deuda {i+1} - Monto Capital en {moneda_global}", min_value=0.0, step=1000.0, value=10000.0, key=f"cap_{i}", help="El saldo actual vigente de este préstamo comercial que la empresa pidió al banco.")
                t_deuda = cols[1].number_input(f"Deuda {i+1} - Tasa Deuda Asignada (% Anual)", min_value=0.0, step=0.1, value=6.0, key=f"td_{i}", help="La tasa de interés anual que el banco cobró a la empresa exactamente al adquirir esta deuda.")
                t_ref = cols[2].number_input(f"Deuda {i+1} - Tasa Ref. Histórica (% Anual)", min_value=0.0, step=0.1, value=4.0, key=f"tref_{i}", help="La Tasa Base Libre de Riesgo (Bono Central o Treasury) que existía a nivel macroeconómico exactamente en la fecha histórica en que se firmó este crédito. El sistema hace una resta para aislar el sobrecargo exacto (Spread por Riesgo de Empresa) que el banco aplicó en ese momento.")
                
                spread = t_deuda - t_ref
                suma_capital += cap
                suma_ponderada += (spread * cap)
                deudas.append({"capital": cap, "tasa_deuda": t_deuda, "tasa_ref_hist": t_ref, "spread": spread})
                
        st.markdown("---")
        
        if st.button("▶ Ejecutar Cálculo de Tasa IBR", type="primary"):
            spread_ponderado = 0.0
            if suma_capital > 0:
                spread_ponderado = suma_ponderada / suma_capital
                
            ibr_final = tasa_base + spread_ponderado
            
            st.subheader("Resultados y Memoria de Cálculo (IBR)")
            st.success(f"**Tasa IBR Anual Final (Tasa Base + Spread): {ibr_final:.4f}%**")
            st.latex(r"Tasa\ IBR = Tasa\ Libre\ Riesgo + \frac{\sum_{j=1}^{m} Spread_j \times Capital_j}{\sum_{j=1}^{m} Capital_j}")
            st.write(f"- Plazo del Contrato: {plazo_meses_ibr} meses ({plazo_anios:.1f} años)")
            st.write(f"- Tasa Base Libre de Riesgo Aplicada ({desc_tasa.replace('Ingrese ', '').replace(' (%)', '')}): {tasa_base:.4f}%")
            st.write(f"- Spread de Riesgo de la Empresa (Ponderado): {spread_ponderado:.4f}%")
            
            import docx
            from docx.shared import Pt, Inches
            import io
            import urllib.request
            import urllib.parse
            
            doc = docx.Document()
            doc.add_heading("MEMORIA DE CÁLCULO - TASA DE DESCUENTO IBR (IFRS 16)", 0)
            doc.add_paragraph(f"Fecha: {datetime.datetime.now().strftime('%Y-%m-%d %H:%M:%S')}")
            
            doc.add_heading("1. Variables Ingresadas", level=1)
            doc.add_paragraph(f"Moneda del Contrato: {moneda_global}")
            doc.add_paragraph(f"Plazo del Contrato: {plazo_meses_ibr} meses ({plazo_anios:.1f} años)")
            doc.add_paragraph(f"Tasa Base Libre de Riesgo Aplicada ({desc_tasa.replace('Ingrese ', '').replace(' (%)', '')}): {tasa_base:.4f}%")
            
            doc.add_heading("2. Detalle de Deudas y Spread", level=1)
            if num_deudas == 0:
                doc.add_paragraph("No se ingresaron deudas referenciales. Spread = 0.0000%")
            else:
                doc.add_paragraph(f"Capital Total Histórico Vigente Ponderado: {suma_capital:,.2f}")
                for i, d in enumerate(deudas):
                    doc.add_paragraph(f"Deuda {i+1}: Capital {d['capital']:,.2f} | Tasa Deuda {d['tasa_deuda']:.4f}% | Ref. Histórica {d['tasa_ref_hist']:.4f}% | Spread {d['spread']:.4f}%")
                doc.add_paragraph(f"RESULTADO SPREAD PONDERADO: {spread_ponderado:.4f}%")
                
            doc.add_heading("3. Fórmula de Cálculo", level=1)
            try:
                formula_ibr = r"\bg_white\dpi{150}\large Tasa\ IBR\ =\ Tasa\ Libre\ Riesgo\ +\ \frac{\sum_{j=1}^{m} Spread_j\ \times\ Capital_j}{\sum_{j=1}^{m} Capital_j}"
                url_ibr = "https://latex.codecogs.com/png.latex?" + urllib.parse.quote(formula_ibr)
                req_ibr = urllib.request.Request(url_ibr, headers={'User-Agent': 'Mozilla/5.0'})
                with urllib.request.urlopen(req_ibr) as response_ibr:
                    img_data_ibr = response_ibr.read()
                doc.add_picture(io.BytesIO(img_data_ibr), width=Inches(4.5))
            except Exception:
                doc.add_paragraph("Tasa IBR = Tasa Base Actual + Spread Ponderado")
            
            doc.add_heading("4. Resultado Final", level=1)
            p = doc.add_paragraph()
            r = p.add_run(f"TASA IBR ANUAL FINAL APLICABLE AL CONTRATO: {ibr_final:.4f}%")
            r.bold = True
            r.font.size = Pt(14)
            
            output = io.BytesIO()
            doc.save(output)
            word_data = output.getvalue()
            
            st.download_button(
                label="📄 Descargar Memoria de Cálculo IBR (Word)",
                data=word_data,
                file_name=f"Reporte_IBR_{datetime.datetime.now().strftime('%Y%m%d')}.docx",
                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
            )
        
    with tab_men:
        st.subheader("Conversor a Tasa Mensual Equivalente")
        st.info("Utilice esta base matemática para transformar la Tasa Anual Efectiva (sea Implícita o IBR) en la tasa periódica exacta que debe imputar en sus asientos mensuales.")
        
        c_anual, c_res = st.columns(2)
        t_anual_input = c_anual.number_input("Tasa Anual Efectiva (%)", min_value=0.0000, max_value=2000.0, value=5.0000, step=0.1000, format="%.4f")
        
        st.markdown("---")
        
        if st.button("▶ Ejecutar Conversión a Tasa Mensual", type="primary"):
            if t_anual_input >= 0:
                t_mensual_calc = (((1 + (t_anual_input / 100.0)) ** (1/12)) - 1) * 100.0
                c_res.success(f"**Tasa Mensual Equivalente:**\n### {t_mensual_calc:.6f}%")
                
                st.latex(r"Tasa\ Mensual\ =\ \left( \sqrt[12]{1\ +\ Tasa\ Anual} \right)\ -\ 1")
                
                st.markdown(f"**¿Cómo contabilizar?**\nAplique la tasa del **{t_mensual_calc:.6f}%** mes a mes sobre el saldo insoluto del pasivo en su cuadro de amortización para reconocer el devengo del gasto financiero mensual bajo IFRS 16.")
                
                import docx
                from docx.shared import Pt, Inches
                import io
                import urllib.request
                import urllib.parse
                
                doc = docx.Document()
                doc.add_heading("MEMORIA DE CÁLCULO - TASA MENSUAL EQUIVALENTE (IFRS 16)", 0)
                doc.add_paragraph(f"Fecha: {datetime.datetime.now().strftime('%Y-%m-%d %H:%M:%S')}")
                
                doc.add_heading("1. Variables Ingresadas", level=1)
                doc.add_paragraph(f"Tasa Anual Efectiva (Implícita o IBR): {t_anual_input:.4f}%")
                
                doc.add_heading("2. Fórmula de Cálculo", level=1)
                doc.add_paragraph("La transformación a base mensual se realiza mediante la fórmula de interés compuesto para mantener la equivalencia temporal del valor del dinero.")
                try:
                    formula_mens = r"\bg_white\dpi{150}\large Tasa\ Mensual\ =\ \left( \sqrt[12]{1\ +\ Tasa\ Anual} \right)\ -\ 1"
                    url_mens = "https://latex.codecogs.com/png.latex?" + urllib.parse.quote(formula_mens)
                    req_mens = urllib.request.Request(url_mens, headers={'User-Agent': 'Mozilla/5.0'})
                    with urllib.request.urlopen(req_mens) as response_mens:
                        img_data_mens = response_mens.read()
                    doc.add_picture(io.BytesIO(img_data_mens), width=Inches(3.5))
                except Exception:
                    doc.add_paragraph("Tasa Mensual = ( (1 + Tasa Anual) ^ (1/12) ) - 1")
                
                doc.add_heading("3. Resultado Final y Aplicación Contable", level=1)
                p = doc.add_paragraph()
                r = p.add_run(f"TASA MENSUAL EQUIVALENTE: {t_mensual_calc:.6f}%")
                r.bold = True
                r.font.size = Pt(14)
                
                doc.add_paragraph(f"Aplicación: Multiplique el saldo insoluto del pasivo por arrendamiento por la tasa mensual de {t_mensual_calc:.6f}% en cada cierre contable para reconocer el gasto financiero devengado en el mes bajo la normativa NIIF 16.")
                
                output = io.BytesIO()
                doc.save(output)
                word_data = output.getvalue()
                
                st.download_button(
                    label="📄 Descargar Memoria de Tasa Mensual (Word)",
                    data=word_data,
                    file_name=f"Reporte_TasaMensual_{datetime.datetime.now().strftime('%Y%m%d')}.docx",
                    mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
                )

    with tab_vp:
        st.subheader("Calculadora de Valor Presente (VP)")
        st.info("💡 **Simulador de Contrato:** Esta herramienta utiliza el motor financiero exacto del sistema IFRS 16 Pro para calcular el monto inicial del Pasivo (VP) y proyectar el Cuadro de Amortización, aplicando descuento de flujos efectivos.")
        
        c_vp1, c_vp2, c_vp3 = st.columns(3)
        vp_canon = c_vp1.number_input("Monto de Cuota (Canon)", min_value=0.0, step=10.0, value=100.0, format="%.4f")
        vp_f_inicio = c_vp2.date_input("Fecha de Inicio", value=date.today(), min_value=date(1900, 1, 1), max_value=date(2100, 12, 31))
        vp_f_fin = c_vp3.date_input("Fecha de Término", value=date.today() + relativedelta(years=1), min_value=date(1900, 1, 1), max_value=date(2100, 12, 31))
        
        c_vp4, c_vp5, c_vp6 = st.columns(3)
        vp_tasa = c_vp4.number_input("Tasa Efectiva Anual (%)", min_value=0.0, step=0.1, value=6.0, format="%.4f")
        
        frecuencias_raw = obtener_parametros('FRECUENCIA_PAGO')
        if not frecuencias_raw: frecuencias_raw = ["Mensual|1", "Trimestral|3", "Semestral|6", "Anual|12"]
        nombres_frec = [f.split('-')[0] for f in frecuencias_raw]
        vp_frec = c_vp5.selectbox("Frecuencia de Pago", nombres_frec, key="vp_frec_input_tab")
        vp_tipo = c_vp6.selectbox("Tipo de Pago", ["Vencido", "Anticipado"], key="vp_tipo_input_tab")
        
        monedas_activas = obtener_parametros('MONEDA')
        if not monedas_activas: monedas_activas = ["UF", "CLP", "USD", "EUR"]
        vp_moneda = st.selectbox("Moneda del Contrato", monedas_activas, key="vp_moneda_tab")
        
        st.markdown("---")
        
        if st.button("▶ Calcular Valor Presente y Tabla de Amortización", type="primary"):
            # 1. Calculo automático de plazo
            diff = relativedelta(vp_f_fin, vp_f_inicio)
            vp_plazo_calc = diff.years * 12 + diff.months
            if diff.days >= 15: vp_plazo_calc += 1
            
            map_frec = {'Mensual': 1}
            for fr in frecuencias_raw:
                parts = fr.split('-')
                if len(parts) == 2 and parts[1].strip().isdigit():
                    map_frec[parts[0].strip()] = int(parts[1].strip())
            
            f_meses = map_frec.get(vp_frec.strip(), 1)
            vp_canon_real = vp_canon # El sistema guarda el Monto por Cuota, no el mensual.
            vp_tasa_mensual = (((1 + (float(vp_tasa) / 100.0)) ** (1/12)) - 1)
            
            st.info(f"ℹ️ **Datos Matemáticos Internos:** La cuota pagadera cada {f_meses} mes(es) es de **{vp_canon_real:,.2f}**, la Tasa Efectiva Mensual aplicada será del **{(vp_tasa_mensual*100):.6f}%**.")

            if vp_plazo_calc <= 0:
                st.error("❌ La Fecha de Término debe ser posterior a la Fecha de Inicio.")
            else:
                # 2. Vincular valor de moneda a la fecha inicial
                tc_inicio = obtener_tc_cache(vp_moneda, vp_f_inicio) if vp_moneda != "CLP" else 1.0
                if tc_inicio <= 0: tc_inicio = 1.0
                
                c_fake = {
                    'Codigo_Interno': 'SIM-VP',
                    'Empresa': 'Simulador',
                    'Clase_Activo': 'Simulador',
                    'Nombre': 'Calculadora VP',
                    'Moneda': vp_moneda,
                    'Valor_Moneda_Inicio': tc_inicio, # Linked to Start Date
                    'Canon': float(vp_canon_real),
                    'Tasa': float(vp_tasa) / 100.0,
                    'Tasa_Mensual': vp_tasa_mensual,
                    'Plazo': int(vp_plazo_calc),
                    'Inicio': vp_f_inicio.strftime('%Y-%m-%d'),
                    'Fin': vp_f_fin.strftime('%Y-%m-%d'),
                    'Estado': 'Activo',
                    'Tipo_Pago': vp_tipo,
                    'Frecuencia_Pago': vp_frec,
                    'Costos_Directos': 0.0,
                    'Pagos_Anticipados': 0.0,
                    'Costos_Desmantelamiento': 0.0,
                    'Incentivos': 0.0,
                    'Ajuste_ROU': 0.0
                }
                
                # 3. Invocar al motor financiero real
                tab_vp_calc, vp_val, rou_val = obtener_motor_financiero(c_fake)
                
                st.success(f"### Valor Presente Calculado (VP): {vp_val:,.2f} {vp_moneda}")
                if vp_moneda != "CLP":
                    st.info(f"💡 Tipo de cambio {vp_moneda} capturado al {vp_f_inicio.strftime('%d-%m-%Y')}: {tc_inicio:,.2f} CLP | VP Equivalente inicial: {(vp_val*tc_inicio):,.0f} CLP")
                
                if not tab_vp_calc.empty:
                    st.write(f"**Cuadro de Amortización Teórico ({vp_plazo_calc} meses, Pago {vp_tipo})**")
                    st.dataframe(tab_vp_calc.style.format(precision=2, thousands="."), use_container_width=True)
                    
                    import io
                    output_xl = io.BytesIO()
                    with pd.ExcelWriter(output_xl, engine='xlsxwriter') as writer:
                        tab_vp_calc.to_excel(writer, index=False)
                    excel_data = output_xl.getvalue()
                    
                    st.download_button(
                        label="📥 Descargar Cuadro a Excel",
                        data=excel_data,
                        file_name="Simulacion_VP_IFRS16.xlsx",
                        mime="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
                        key="down_vp_sim"
                    )
                    
                    # 4. Generar Memoria Word
                    import docx
                    from docx.shared import Pt, Inches
                    import io
                    import urllib.request
                    import urllib.parse
                    
                    doc = docx.Document()
                    doc.add_heading("MEMORIA DE CÁLCULO - VALOR PRESENTE (IFRS 16)", 0)
                    doc.add_paragraph(f"Fecha de cálculo: {datetime.datetime.now().strftime('%Y-%m-%d %H:%M:%S')}")
                    
                    doc.add_heading("1. Variables de Entrada", level=1)
                    doc.add_paragraph(f"Moneda del Contrato: {vp_moneda}")
                    if vp_moneda != "CLP": doc.add_paragraph(f"Tipo de Cambio inicial (UF/USD/EUR): {tc_inicio:,.2f} al {vp_f_inicio.strftime('%Y-%m-%d')}")
                    doc.add_paragraph(f"Canon de Arriendo Mensual: {vp_canon:,.2f} {vp_moneda}")
                    doc.add_paragraph(f"Fecha de Inicio: {vp_f_inicio.strftime('%Y-%m-%d')}")
                    doc.add_paragraph(f"Fecha de Término: {vp_f_fin.strftime('%Y-%m-%d')}")
                    doc.add_paragraph(f"Plazo Calculado para Cuotas: {vp_plazo_calc} meses")
                    doc.add_paragraph(f"Tasa Efectiva Anual (Descuento): {vp_tasa:.4f}%")
                    doc.add_paragraph(f"Tipo de Pago: {vp_tipo}")
                    doc.add_paragraph(f"Frecuencia de Pago Real: {vp_frec}")
                    
                    doc.add_heading("2. Fórmula Aplicada", level=1)
                    doc.add_paragraph("Para proyectar el Cuadro de Amortización se emplea la siguiente fórmula estándar NIIF 16 que descuenta los flujos futuros. Si el pago es anticipado, el primer flujo ocurre en el mes cero (sin descuento).")
                    try:
                        formula_vp = r"\bg_white\dpi{150}\large VP\ =\ \sum_{t=x}^{n}\ \frac{Cuota}{(1+i_{mensual})^t}"
                        url_vp = "https://latex.codecogs.com/png.latex?" + urllib.parse.quote(formula_vp)
                        req_vp = urllib.request.Request(url_vp, headers={'User-Agent': 'Mozilla/5.0'})
                        with urllib.request.urlopen(req_vp) as response_vp:
                            img_data_vp = response_vp.read()
                        doc.add_picture(io.BytesIO(img_data_vp), width=Inches(3.5))
                    except Exception:
                        doc.add_paragraph("Fórmula Módica: Sumatoria(Cuota / (1 + Tasa_Mensual)^t)")
                    doc.add_paragraph("Donde 'x=0' para opciones de Pago Anticipado, y 'x=1' para Pago Vencido.")
                    
                    doc.add_heading("3. Resultado Final", level=1)
                    p = doc.add_paragraph()
                    r = p.add_run(f"VALOR PRESENTE (VP): {vp_val:,.2f} {vp_moneda}")
                    r.bold = True
                    r.font.size = Pt(14)
                    
                    if vp_moneda != "CLP":
                        doc.add_paragraph(f"Equivalente referencial inicial: {(vp_val*tc_inicio):,.0f} CLP (usando T.C. inicial {tc_inicio:,.2f})")
                        
                    output_vp = io.BytesIO()
                    doc.save(output_vp)
                    word_data_vp = output_vp.getvalue()
                    
                    st.download_button(
                        label="📄 Descargar Memoria de Cálculo VP (Word)",
                        data=word_data_vp,
                        file_name=f"Reporte_ValorPresente_{datetime.datetime.now().strftime('%Y%m%d')}.docx",
                        mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                        key="down_vp_word_sim"
                    )
                else:
                    st.error("Los parámetros ingresados no generaron un cuadro válido.")

def main():
    inicializar_db() # Garantizar que la BD exista al iniciar
    if not st.session_state.auth:
        html_login = """
        <div style="display: flex; flex-direction: column; align-items: flex-start; justify-content: center; margin-bottom: 25px; margin-top: 15px;">
            <div style="display: flex; align-items: baseline; font-family: 'Montserrat', sans-serif;">
                <span style="font-size: 1.6rem; margin-right: 8px;">🔐</span>
                <span style="color: #00829B; font-weight: 500; font-size: 1.4rem; letter-spacing: 2px; margin-right: 18px;">LOGIN</span>
                <span style="color: #00829B; font-weight: 900; font-size: 2.4rem; letter-spacing: -2px;">mund</span>
                <div style="position: relative; display: inline-block; margin-left: 2px; top: 1px;">
                    <div style="box-sizing: content-box; width: 13px; height: 13px; border: 6px solid #00829B; border-radius: 50%; background-color: #FFCE00; display: inline-block;"></div>
                    <div style="position: absolute; top: -5px; left: -5px; width: 7px; height: 7px; background-color: #FF8A00; border-radius: 50%;"></div>
                    <div style="position: absolute; top: -11px; right: 0px; width: 5px; height: 5px; background-color: #E6007E; border-radius: 50%;"></div>
                    <div style="position: absolute; top: -2px; right: -6px; width: 4px; height: 4px; background-color: #8C2482; border-radius: 50%;"></div>
                    <div style="position: absolute; top: 9px; right: -10px; width: 7px; height: 7px; background-color: #E3000F; border-radius: 50%;"></div>
                    <div style="position: absolute; top: 20px; right: -5px; width: 6px; height: 6px; background-color: #00B4E6; border-radius: 50%;"></div>
                </div>
                <span style="color: #00829B; font-weight: 900; font-size: 2.4rem; margin-left: 10px; letter-spacing: -2px;">16</span>
            </div>
            <div style="width: 100%; max-width: 380px; height: 3px; background-color: #FFCE00; border-radius: 10px; margin-top: 8px;"></div>
        </div>
        """
        st.markdown(html_login, unsafe_allow_html=True)
        u = st.text_input("Usuario", "admin")
        p = st.text_input("Contraseña", type="password")
        if st.button("Entrar"):
            rol = verificar_credenciales(u, p)
            if rol:
                st.session_state.auth = True
                st.session_state.rol = rol
                st.session_state.user = u
                st.rerun()
            else:
                st.error("❌ Credenciales incorrectas")
    else:
        import time
        t0 = time.time()
        col_espacio, col_reloj = st.columns([8, 2])
        with col_reloj:
            reloj_ui = st.empty()
            
        html_logo = """
        <div style="display: flex; flex-direction: column; align-items: flex-start; margin-bottom: 20px;">
            <div style="display: flex; align-items: baseline; font-family: 'Montserrat', sans-serif;">
                <span style="color: #00829B; font-weight: 900; font-size: 2.8rem; letter-spacing: -2px;">mund</span>
                <div style="position: relative; display: inline-block; margin-left: 2px; top: 2px;">
                    <div style="box-sizing: content-box; width: 15px; height: 15px; border: 8px solid #00829B; border-radius: 50%; background-color: #FFCE00; display: inline-block;"></div>
                    <div style="position: absolute; top: -6px; left: -6px; width: 8px; height: 8px; background-color: #FF8A00; border-radius: 50%;"></div>
                    <div style="position: absolute; top: -14px; right: 0px; width: 6px; height: 6px; background-color: #E6007E; border-radius: 50%;"></div>
                    <div style="position: absolute; top: -2px; right: -8px; width: 5px; height: 5px; background-color: #8C2482; border-radius: 50%;"></div>
                    <div style="position: absolute; top: 10px; right: -12px; width: 9px; height: 9px; background-color: #E3000F; border-radius: 50%;"></div>
                    <div style="position: absolute; top: 22px; right: -6px; width: 8px; height: 8px; background-color: #00B4E6; border-radius: 50%;"></div>
                </div>
                <span style="color: #00829B; font-weight: 900; font-size: 2.8rem; margin-left: 12px; letter-spacing: -2px;">16</span>
            </div>
            <div style="color: #00829B; font-size: 0.6rem; font-weight: 800; letter-spacing: 2px; margin-top: -5px; padding-left: 5px; font-family: 'Montserrat', sans-serif;">
                FIBRA | MÓVIL | TV | FIJO
            </div>
        </div>
        """
        st.sidebar.markdown(html_logo, unsafe_allow_html=True)
        st.sidebar.markdown("---")
        
        rol_actual = st.session_state.get('rol', 'Lector')
        usuario_actual = st.session_state.get('user', 'admin')
        st.sidebar.info(f"👤 **{usuario_actual}**\n\n🛡️ Rol: {rol_actual}")
        
        st.sidebar.button("Salir (Cerrar Sesión)", on_click=lambda: st.session_state.clear() or st.session_state.update(auth=False))
        
        # --- DEFINICIÓN DE MENÚ RBAC ---
        menus_todas = ["Monedas", "Contratos", "Resumen de Saldos", "Asientos", "Nota: Movimiento de saldos", "Nota: Vencimientos NIIF 16", "Auditoría", "Asistente de calculos (tasas de contratos-Activo y pasivo ROU)", "Configuración"]
        
        if rol_actual == 'Administrador':
            opciones_menu = menus_todas
        elif rol_actual == 'Analista Financiero (Editor)':
            opciones_menu = [m for m in menus_todas if m not in ['Configuración', 'Auditoría']]
        elif rol_actual == 'Auditor Ext. / Gerencia (Lector)':
            opciones_menu = ["Monedas", "Contratos", "Resumen de Saldos", "Asientos", "Nota: Movimiento de saldos", "Nota: Vencimientos NIIF 16", "Auditoría", "Asistente de calculos (tasas de contratos-Activo y pasivo ROU)"]
        elif rol_actual == 'Ingeniero IT (Técnico)':
            opciones_menu = ["Configuración"]
        else:
            opciones_menu = ["Resumen de Saldos"]

        op = st.sidebar.radio("Menú Principal", opciones_menu)
        if op == "Monedas": modulo_monedas()
        elif op == "Contratos": modulo_contratos()
        elif op == "Resumen de Saldos": modulo_dashboard()
        elif op == "Asientos": modulo_asientos()
        elif op == "Nota: Movimiento de saldos": modulo_notas()
        elif op == "Nota: Vencimientos NIIF 16": modulo_vencimientos()
        elif op == "Auditoría": modulo_auditoria()
        elif op == "Asistente de calculos (tasas de contratos-Activo y pasivo ROU)": modulo_asistente_ibr()
        elif op == "Configuración": modulo_configuracion()
        
        t1 = time.time()
        reloj_ui.markdown(f"<div style='text-align: right; color: gray; font-size: 0.9em; margin-top: -40px;'>⏱️ Tiempo de ejecución: {t1 - t0:.2f} s</div>", unsafe_allow_html=True)

if __name__ == "__main__": main()